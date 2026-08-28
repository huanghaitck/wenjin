from __future__ import annotations

import argparse
import hashlib
import json
import re
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document


def extract_text(path: Path) -> str:
    document = Document(path)
    blocks = [paragraph.text for paragraph in document.paragraphs]
    blocks.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(block for block in blocks if block.strip())


def normalized(text: str) -> str:
    return re.sub(r"[^0-9A-Za-zА-Яа-яЁё\u4e00-\u9fff]+", "", text).casefold()


def shingles(text: str, width: int = 17) -> set[str]:
    if len(text) <= width:
        return {text} if text else set()
    return {hashlib.sha1(text[index:index + width].encode("utf-8")).hexdigest()[:12] for index in range(0, len(text) - width + 1, 4)}


def similarity(left: dict, right: dict) -> tuple[float, float]:
    common = len(left["shingles"] & right["shingles"])
    union = len(left["shingles"] | right["shingles"])
    smaller = min(len(left["shingles"]), len(right["shingles"]))
    return (common / union if union else 0.0, common / smaller if smaller else 0.0)


def audit(paths: list[Path]) -> dict:
    records = []
    for path in sorted({path.resolve() for path in paths if path.suffix.lower() == ".docx" and path.is_file()}):
        if path.name.startswith("~$"):
            continue
        try:
            text = extract_text(path)
        except Exception as exc:
            records.append({"path": str(path), "error": f"{type(exc).__name__}: {exc}"})
            continue
        compact = normalized(text)
        stat = path.stat()
        records.append({
            "path": str(path), "name": path.name, "modified_ns": stat.st_mtime_ns,
            "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(timespec="seconds"),
            "characters": len(compact), "text_sha256": hashlib.sha256(compact.encode("utf-8")).hexdigest(),
            "first_line": next((line.strip() for line in text.splitlines() if line.strip()), "")[:240],
            "shingles": shingles(compact),
        })
    valid = [record for record in records if "error" not in record]
    parent = list(range(len(valid)))

    def find(index: int) -> int:
        while parent[index] != index:
            parent[index] = parent[parent[index]]; index = parent[index]
        return index

    def union(left: int, right: int) -> None:
        left, right = find(left), find(right)
        if left != right:
            parent[right] = left

    pair_receipts = {}
    for left in range(len(valid)):
        for right in range(left + 1, len(valid)):
            if valid[left]["text_sha256"] == valid[right]["text_sha256"]:
                union(left, right); pair_receipts[(left, right)] = (1.0, 1.0); continue
            jaccard, containment = similarity(valid[left], valid[right])
            if jaccard >= 0.68 or containment >= 0.90:
                union(left, right); pair_receipts[(left, right)] = (jaccard, containment)
    clustered = defaultdict(list)
    for index in range(len(valid)):
        clustered[find(index)].append(index)
    groups = []
    for indices in clustered.values():
        if len(indices) < 2:
            continue
        longest = max(valid[index]["characters"] for index in indices)
        complete = [index for index in indices if valid[index]["characters"] >= longest * 0.95]
        keeper = max(complete, key=lambda index: (valid[index]["modified_ns"], valid[index]["characters"]))
        members = []
        for index in sorted(indices, key=lambda value: valid[value]["modified_ns"]):
            jaccard, containment = (1.0, 1.0) if index == keeper else similarity(valid[keeper], valid[index])
            members.append({key: value for key, value in valid[index].items() if key != "shingles"} | {
                "recommended_current": index == keeper,
                "jaccard_to_current": round(jaccard, 4),
                "containment_to_current": round(containment, 4),
            })
        groups.append({"member_count": len(members), "recommended_current": valid[keeper]["path"], "members": members})
    return {
        "file_count": len(records), "readable_count": len(valid), "group_count": len(groups),
        "groups": sorted(groups, key=lambda group: (-group["member_count"], group["recommended_current"])),
        "errors": [record for record in records if "error" in record],
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Content-based DOCX version clustering; never modifies source files.")
    parser.add_argument("paths", nargs="+", type=Path)
    args = parser.parse_args()
    expanded = []
    for path in args.paths:
        expanded.extend(path.rglob("*.docx") if path.is_dir() else [path])
    print(json.dumps(audit(expanded), ensure_ascii=False, indent=2, default=list))


if __name__ == "__main__":
    main()
