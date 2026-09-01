from __future__ import annotations

import io
import json
import os
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from research_workbench.authoring import import_manuscript
from research_workbench.desktop_runtime import bootstrap_desktop
from research_workbench.document_model import ensure_document, reimport_docx
from research_workbench.model_settings import discover_models, public_settings, reasoning_controls, save_role
from research_workbench.service import initialize_project
from research_workbench.web import build_server
from research_workbench.agent_runtime import list_threads


class D5DesktopPackagingTests(unittest.TestCase):
    def test_reasoning_controls_follow_the_selected_model(self) -> None:
        self.assertEqual(reasoning_controls("openai_compatible", "deepseek-v4-flash", "https://api.deepseek.com")["efforts"], ["low", "high", "max"])
        self.assertEqual(reasoning_controls("openai_compatible", "deepseek-v4-pro", "https://api.deepseek.com")["efforts"], ["high", "max"])
        self.assertEqual(reasoning_controls("ollama", "gpt-oss:20b")["modes"], ["deep"])
        self.assertEqual(reasoning_controls("openai_compatible", "unknown-model")["modes"], [])

    def test_serve_infers_desktop_roots_from_a_registered_project(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            paths = bootstrap_desktop(Path(directory))
            server = build_server(Path(paths["project_root"]), port=0)
            try:
                self.assertEqual(server.workspace_root, Path(paths["workspace_root"]).resolve())
                self.assertEqual(server.library_root, Path(paths["library_root"]).resolve())
                self.assertEqual(server.config_root, Path(paths["config_root"]).resolve())
            finally:
                server.server_close()

    def test_first_run_creates_stable_local_workspace(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            first = bootstrap_desktop(Path(directory))
            second = bootstrap_desktop(Path(directory))
            self.assertEqual(first["project_root"], second["project_root"])
            self.assertTrue((Path(first["project_root"]) / "project.sqlite3").is_file())
            self.assertTrue((Path(first["workspace_root"]) / "workspace.json").is_file())
            self.assertTrue(Path(first["library_root"]).is_dir())
            self.assertEqual(len(list_threads(Path(first["project_root"]))), 1)
            self.assertEqual(list_threads(Path(first["project_root"]))[0]["title"], "新的研究讨论")

    def test_desktop_recovers_an_existing_unregistered_project_before_empty_default(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            data_root = Path(directory)
            existing = data_root / "workspace" / "projects" / "existing-history"
            initialize_project(existing, "已有历史项目")
            source = data_root / "source.pdf"
            source.write_bytes(b"existing source")
            from research_workbench.service import register_source
            register_source(existing, source, "已有材料")
            boot = bootstrap_desktop(data_root)
            self.assertEqual(Path(boot["project_root"]), existing.resolve())
            registry = (data_root / "workspace" / "workspace.json").read_text(encoding="utf-8")
            self.assertIn("已有历史项目", registry)
            backups = list((data_root / "backups" / "projects").glob("*/manifest.json"))
            self.assertEqual(len(backups), 1)

    def test_ollama_role_settings_are_applied_without_a_secret_file(self) -> None:
        with tempfile.TemporaryDirectory() as directory, patch.dict(os.environ, {}, clear=False):
            root = Path(directory)
            result = save_role(root, "main_reasoning", {
                "provider": "ollama", "model": "local-history-model",
                "base_url": "http://127.0.0.1:11434", "timeout_seconds": 120,
            })
            text = (root / "model-settings.json").read_text(encoding="utf-8")
            self.assertNotIn("api_key", text.lower())
            self.assertEqual(os.environ["HRW_AGENT_MODEL"], "local-history-model")
            self.assertEqual(os.environ["HRW_AGENT_BASE_URL"], "http://127.0.0.1:11434")
            self.assertEqual(len(result["roles"]), 9)
            expected = "windows_credential_manager" if os.name == "nt" else "macos_keychain" if sys.platform == "darwin" else "unavailable"
            self.assertEqual(public_settings(root)["credential_backend"], expected)

    def test_model_settings_distinguish_direct_and_reserved_auxiliary_routes(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            roles = {item["role"]: item for item in public_settings(Path(directory))["roles"]}
        self.assertEqual(roles["main_reasoning"]["model"], "deepseek-v4-flash")
        self.assertEqual(roles["main_reasoning"]["preset_id"], "deepseek")
        self.assertEqual(roles["main_reasoning"]["provider"], "disabled")
        self.assertTrue(roles["vision_ocr"]["direct_route"])
        self.assertTrue(roles["domain_agent"]["direct_route"])
        self.assertTrue(roles["translation_helper"]["direct_route"])
        self.assertTrue(roles["review_secondary"]["direct_route"])
        self.assertTrue(roles["web_research"]["direct_route"])
        self.assertTrue(roles["context_compression"]["direct_route"])
        self.assertTrue(roles["title_generation"]["direct_route"])

    def test_word_reimport_creates_a_new_revision_without_overwriting(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            project = Path(directory) / "project"
            initialize_project(project, "Word 往返测试")
            manuscript = import_manuscript(project, "测试稿", "# 导言\n\n原始段落。")
            before = ensure_document(project, manuscript["manuscript_id"])
            word = Document()
            word.add_heading("导言", level=1)
            word.add_paragraph("在 Microsoft Word 中人工修改后的段落。")
            table = word.add_table(rows=2, cols=2)
            table.cell(0, 0).text, table.cell(0, 1).text = "比较项", "数量"
            table.cell(1, 0).text, table.cell(1, 1).text = "停驻点", "12"
            payload = io.BytesIO()
            word.save(payload)
            after = reimport_docx(project, manuscript["manuscript_id"], payload.getvalue())
            self.assertNotEqual(before["current_revision_id"], after["current_revision_id"])
            self.assertEqual(after["revisions"][0]["source_format"], "docx_reimport")
            self.assertIn("人工修改", after["document"]["children"][0]["children"][0]["text"])
            self.assertEqual(after["document"]["children"][0]["children"][1]["rows"][1], ["停驻点", "12"])
            self.assertGreaterEqual(len(after["revisions"]), 2)

    def test_ui_exposes_desktop_word_and_model_bridges(self) -> None:
        root = Path(__file__).parents[1] / "src" / "research_workbench" / "web_assets"
        html = (root / "index.html").read_text(encoding="utf-8")
        script = (root / "app.js").read_text(encoding="utf-8")
        for identifier in ("chooseFolder", "choosePdf", "chooseDocx", "openInWord", "reimportWord"):
            self.assertIn(f'id="{identifier}"', html)
        self.assertIn("/api/model-settings/save", script)
        self.assertIn("/api/model-settings/models", script)
        self.assertIn("Refresh models", script)
        self.assertIn("Windows 凭据管理器", script)

    def test_model_discovery_normalizes_ollama_and_openai_compatible_lists(self) -> None:
        class Response:
            status = 200
            def __init__(self, payload): self.payload = payload
            def __enter__(self): return self
            def __exit__(self, *_): return None
            def read(self): return json.dumps(self.payload).encode("utf-8")
        with tempfile.TemporaryDirectory() as directory:
            config = Path(directory)
            with patch("research_workbench.model_settings.urlopen", return_value=Response({
                "models": [{"name": "qwen3"}, {"name": "deepseek-r1"}, {"name": "qwen3"}],
            })):
                ollama = discover_models(config, "main_reasoning", "ollama", "http://127.0.0.1:11434")
            self.assertEqual(ollama["models"], ["deepseek-r1", "qwen3"])
            with patch("research_workbench.model_settings.urlopen", return_value=Response({
                "data": [{"id": "deepseek-chat"}, {"id": "deepseek-reasoner"}],
            })):
                remote = discover_models(
                    config, "main_reasoning", "openai_compatible", "https://api.example/v1", "secret"
                )
        self.assertEqual(remote["models"], ["deepseek-chat", "deepseek-reasoner"])

    def test_desktop_sidecar_bundles_https_runtime(self) -> None:
        build_script = (Path(__file__).parents[1] / "scripts" / "build_desktop.ps1").read_text(encoding="utf-8")
        self.assertIn("libssl-3-x64.dll", build_script)
        self.assertIn("libcrypto-3-x64.dll", build_script)
        self.assertIn('"--collect-data", "certifi"', build_script)
        tauri_root = Path(__file__).parents[1] / "src-tauri"
        tauri = (tauri_root / "tauri.conf.json").read_text(encoding="utf-8")
        tauri_windows = (tauri_root / "tauri.windows.conf.json").read_text(encoding="utf-8")
        tauri_macos = (tauri_root / "tauri.macos.conf.json").read_text(encoding="utf-8")
        package = (Path(__file__).parents[1] / "package.json").read_text(encoding="utf-8")
        self.assertIn("agent-browser-win32-x64.exe", tauri_windows)
        self.assertIn("agent-browser-APACHE-2.0.txt", tauri)
        self.assertIn('"type": "downloadBootstrapper"', tauri_windows)
        self.assertIn("agent-browser-darwin-arm64", tauri_macos)
        self.assertIn('"signingIdentity": "-"', tauri_macos)
        self.assertIn('"agent-browser": "0.33.0"', package)
        offline_installer = (Path(__file__).parents[1] / "scripts" / "install-wenjin.cmd").read_text(encoding="utf-8")
        self.assertIn("MicrosoftEdgeWebView2RuntimeInstallerX64.exe", offline_installer)
        offline_entry = (Path(__file__).parents[1] / "scripts" / "双击这里离线安装问津.cmd").read_text(encoding="utf-8")
        self.assertIn("install-wenjin.cmd", offline_entry)
        self.assertIn("F3017226-FE2A-4295-8BDF-00C3A9A7E4C5", offline_installer)
        self.assertIn("smoke_desktop_sidecar.py", build_script)

    def test_desktop_failure_screen_has_real_diagnostics(self) -> None:
        root = Path(__file__).parents[1]
        shell = (root / "desktop-shell" / "index.html").read_text(encoding="utf-8")
        script = (root / "desktop-shell" / "desktop.js").read_text(encoding="utf-8")
        rust = (root / "src-tauri" / "src" / "main.rs").read_text(encoding="utf-8")
        self.assertIn("打开启动日志", shell)
        self.assertIn("open_sidecar_log", script)
        self.assertIn("open_sidecar_log", rust)
        self.assertIn("'open_path'", script)
        self.assertIn("fn open_path", rust)
        self.assertIn("CommandEvent::Terminated", rust)

    def test_desktop_shell_does_not_reload_workbench_during_startup(self) -> None:
        script = (Path(__file__).parents[1] / "desktop-shell" / "desktop.js").read_text(encoding="utf-8")
        rust = (Path(__file__).parents[1] / "src-tauri" / "src" / "main.rs").read_text(encoding="utf-8")
        self.assertNotIn("setInterval(() => { frame.src = url; }, 1000)", script)
        self.assertEqual(script.count("frame.src = url"), 1)
        self.assertIn("frameTimeout = window.setTimeout", script)
        self.assertGreater(rust.index("*startup_url = Some"), rust.index("if !ready"))

    def test_desktop_build_python_is_configurable(self) -> None:
        build_script = (Path(__file__).parents[1] / "scripts" / "build_desktop.ps1").read_text(encoding="utf-8")
        rust = (Path(__file__).parents[1] / "src-tauri" / "src" / "main.rs").read_text(encoding="utf-8")
        self.assertIn("HRW_BUILD_PYTHON", build_script)
        self.assertIn("sys.version_info >= (3, 13)", build_script)
        self.assertIn("import sys, docx, fitz, mcp", build_script)
        self.assertIn("No suitable Python 3.13+ build environment", build_script)
        self.assertNotIn('D:\\AI_Workflows\\conda-envs\\historical-research-workbench\\python.exe', build_script)
        self.assertIn("WENJIN_DATA_ROOT", rust)


if __name__ == "__main__":
    unittest.main()
