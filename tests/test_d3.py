from __future__ import annotations

import io
import json
import tempfile
import unittest
from pathlib import Path

from docx import Document

from research_workbench.agent_runtime import create_thread, send_message
from research_workbench.authoring import (
    create_journal_template, ensure_journal_templates, import_manuscript, manuscript_detail,
    save_submission_profile, submission_profiles,
)
from research_workbench.db import connect
from research_workbench.citations import create_note
from research_workbench.document_model import (
    ensure_document,
    export_document,
    import_docx,
    preview_document_export,
    save_document,
)
from research_workbench.service import initialize_project, save_source_citation_metadata


class D3ResearchObjectWorkspaceTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary = tempfile.TemporaryDirectory()
        self.project = Path(self.temporary.name) / "project"
        initialize_project(self.project, "D3 project")
        self.manuscript = import_manuscript(
            self.project, "页面关系论文", "# 导言\n\n1908年，材料称“队伍北行”。[^1]\n\n# 第一节\n\n原有论述。",
        )

    def tearDown(self) -> None:
        self.temporary.cleanup()

    def test_legacy_sections_become_structured_document_without_losing_text(self) -> None:
        detail = ensure_document(self.project, self.manuscript["manuscript_id"])
        self.assertEqual([item["heading"] for item in detail["document"]["children"]], ["导言", "第一节"])
        self.assertIn("1908年", detail["document"]["children"][0]["children"][0]["text"])
        self.assertEqual(detail["source_format"], "legacy_sections")

    def test_save_creates_immutable_revision_and_updates_legacy_read(self) -> None:
        first = ensure_document(self.project, self.manuscript["manuscript_id"])
        tree = first["document"]
        tree["children"][0]["children"][0]["text"] += " 人工补充一句。[EVID:EVT_manual]"
        second = save_document(self.project, self.manuscript["manuscript_id"], tree)
        self.assertNotEqual(first["current_revision_id"], second["current_revision_id"])
        self.assertEqual(len(second["revisions"]), 2)
        with connect(self.project) as connection:
            original = connection.execute(
                "SELECT document_json FROM document_revisions WHERE revision_id = ?",
                (first["current_revision_id"],),
            ).fetchone()[0]
        self.assertNotIn("人工补充一句", original)
        self.assertIn("人工补充一句", manuscript_detail(self.project, self.manuscript["manuscript_id"])["sections"][0]["content"])
        with connect(self.project) as connection:
            counts = [row[0] for row in connection.execute(
                "SELECT COUNT(*) FROM section_versions GROUP BY section_id ORDER BY section_id"
            )]
            evidence_refs = connection.execute(
                "SELECT evidence_refs_json FROM section_versions ORDER BY created_at DESC LIMIT 1"
            ).fetchone()[0]
        self.assertEqual(sorted(counts), [1, 2])
        self.assertEqual(evidence_refs, '["EVT_manual"]')

    def test_context_dialogue_freezes_revision_section_and_selection_hash(self) -> None:
        detail = ensure_document(self.project, self.manuscript["manuscript_id"])
        section = detail["document"]["children"][0]
        node = section["children"][0]
        thread = create_thread(self.project, "稿件讨论")
        result = send_message(self.project, thread["thread_id"], "这句话的证据边界如何？", {
            "manuscript_id": self.manuscript["manuscript_id"],
            "revision_id": detail["current_revision_id"],
            "section_id": section["section_id"],
            "node_id": node["node_id"],
            "selection_text": "队伍北行",
            "attached_refs": [],
        })
        binding = result["messages"][0]["context_binding"]
        self.assertEqual(binding["revision_id"], detail["current_revision_id"])
        self.assertEqual(binding["selection_text"], "队伍北行")
        self.assertEqual(len(binding["selection_hash"]), 64)

    def test_save_can_insert_a_section_before_the_current_first_section(self) -> None:
        detail = ensure_document(self.project, self.manuscript["manuscript_id"])
        tree = detail["document"]
        tree["children"].insert(0, {
            "type": "section", "node_id": "NOD_abstract", "section_id": "",
            "heading": "摘要与关键词",
            "children": [{"type": "paragraph", "node_id": "NOD_abstract_text", "text": "摘要：测试。"}],
        })
        saved = save_document(self.project, self.manuscript["manuscript_id"], tree)
        self.assertEqual(
            [section["heading"] for section in saved["document"]["children"]],
            ["摘要与关键词", "导言", "第一节"],
        )
        self.assertTrue(saved["document"]["children"][0]["section_id"].startswith("SEC_"))

    def test_docx_and_markdown_adapters_return_fidelity_receipts(self) -> None:
        package = io.BytesIO()
        docx = Document()
        docx.add_heading("导言", level=1)
        docx.add_paragraph("正文段落。")
        table = docx.add_table(rows=3, cols=3)
        for column, value in enumerate(("比较项", "谭卫道", "李希霍芬")):
            table.cell(0, column).text = value
        table.cell(1, 0).text, table.cell(1, 1).text, table.cell(1, 2).text = "旅行年份", "1873", "1872"
        table.cell(2, 0).text, table.cell(2, 1).text, table.cell(2, 2).text = "停驻点", "12", "9"
        docx.save(package)
        imported = import_docx(self.project, "DOCX 稿件", package.getvalue())
        self.assertEqual(imported["import_fidelity"]["level"], "limited")
        self.assertEqual(imported["document"]["children"][0]["children"][1]["type"], "table")
        markdown = export_document(self.project, imported["manuscript_id"], "markdown")
        word = export_document(self.project, imported["manuscript_id"], "docx")
        self.assertTrue((self.project / markdown["project_path"]).is_file())
        self.assertTrue((self.project / word["project_path"]).is_file())
        self.assertIn("| 比较项 | 谭卫道 | 李希霍芬 |", (self.project / markdown["project_path"]).read_text(encoding="utf-8"))
        self.assertEqual(len(Document(self.project / word["project_path"]).tables), 1)
        self.assertEqual(word["fidelity"]["level"], "structured_with_true_footnotes")

    def test_export_does_not_repeat_an_empty_title_section(self) -> None:
        manuscript = import_manuscript(
            self.project,
            "同名标题",
            "# 同名标题\n\n# 正文\n\n正文段落。",
        )

        markdown = export_document(self.project, manuscript["manuscript_id"], "markdown")
        markdown_text = (self.project / markdown["project_path"]).read_text(encoding="utf-8")
        self.assertEqual(markdown_text.count("# 同名标题"), 1)

        word = export_document(self.project, manuscript["manuscript_id"], "docx")
        document = Document(self.project / word["project_path"])
        self.assertEqual([paragraph.text for paragraph in document.paragraphs].count("同名标题"), 1)

    def test_docx_export_renders_simple_markdown_bold_without_asterisks(self) -> None:
        manuscript = import_manuscript(
            self.project, "粗体导出", "# 英文信息\n\n**Title:** Example",
        )

        word = export_document(self.project, manuscript["manuscript_id"], "docx")
        document = Document(self.project / word["project_path"])
        paragraph = next(item for item in document.paragraphs if "Title:" in item.text)
        self.assertEqual(paragraph.text, "Title: Example")
        self.assertTrue(any(run.text == "Title:" and run.bold for run in paragraph.runs))

    def test_sequential_reference_export_reuses_source_number_and_reports_missing_gates(self) -> None:
        manuscript = import_manuscript(
            self.project, "顺序编码稿", "# 正文\n\n甲事。[EVID:EVT_one] 乙事。[EVID:EVT_two]",
        )
        with connect(self.project) as connection:
            project_id = connection.execute("SELECT project_id FROM projects").fetchone()[0]
            connection.execute(
                "INSERT INTO sources VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
                ("SRC_book", project_id, "测试书", "local_file", "book.pdf", "acquired", "ready", "partial", "2026-01-01"),
            )
            payload = {"claims": [{"evidence": [
                {"evidence_id": "EVT_one", "source_id": "SRC_book", "printed_pages": ["12"]},
                {"evidence_id": "EVT_two", "source_id": "SRC_book", "printed_pages": ["15"]},
            ]}]}
            connection.execute(
                "INSERT INTO evidence_freezes(freeze_id, title, status, payload_json, approved_by, approved_at, created_at) VALUES (?, ?, ?, ?, ?, ?, ?)",
                ("FRZ_test", "测试冻结", "approved", json.dumps(payload), "tester", "2026-01-01", "2026-01-01"),
            )
        save_source_citation_metadata(self.project, "SRC_book", {
            "author": "张三", "title": "测试书", "place": "西安", "publisher": "测试出版社",
            "year": "1908", "type_code": "M", "verified_by": "tester",
        })
        template = create_journal_template(
            self.project, "《顺序编码测试》", "参考文献置于文后，按正文出现顺序全文连续编号",
            ["正文", "参考文献", "英文摘要", "作者信息"], requirements={},
        )
        exported = export_document(self.project, manuscript["manuscript_id"], "markdown", template["template_id"])
        text = (self.project / exported["project_path"]).read_text(encoding="utf-8")
        self.assertIn("甲事。[1]12", text)
        self.assertIn("乙事。[1]15", text)
        self.assertEqual(text.count("[1] 张三. 测试书[M]. 西安：测试出版社，1908"), 1)
        self.assertEqual(exported["fidelity"]["citation_status"], "BLOCKED")
        self.assertTrue(any("英文题名" in warning for warning in exported["fidelity"]["warnings"]))
        word = export_document(self.project, manuscript["manuscript_id"], "docx", "builtin-tangdu-current")
        import zipfile
        with zipfile.ZipFile(self.project / word["project_path"]) as package:
            document_xml = package.read("word/document.xml").decode("utf-8")
            styles_xml = package.read("word/styles.xml").decode("utf-8")
        self.assertIn('w:vertAlign w:val="superscript"', document_xml)
        self.assertIn("[1]12", document_xml)
        self.assertIn('w:ascii="Times New Roman"', styles_xml)
        self.assertIn('w:eastAsia="宋体"', styles_xml)
        self.assertIn('w:eastAsia="黑体"', styles_xml)

    def test_generated_references_replace_placeholder_section(self) -> None:
        manuscript = import_manuscript(
            self.project, "顺序编码稿",
            "# 正文\n\n甲事。[EVID:EVT_one]\n\n# 参考文献\n\n（由工作台生成）\n\n# 英文摘要\n\nEnglish",
        )
        with connect(self.project) as connection:
            project_id = connection.execute("SELECT project_id FROM projects").fetchone()[0]
            connection.execute(
                "INSERT INTO sources VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
                ("SRC_book", project_id, "测试书", "local_file", "book.pdf", "acquired", "ready", "partial", "2026-01-01"),
            )
            payload = {"claims": [{"evidence": [
                {"evidence_id": "EVT_one", "source_id": "SRC_book", "printed_pages": ["12"]},
            ]}]}
            connection.execute(
                "INSERT INTO evidence_freezes(freeze_id, title, status, payload_json, approved_by, approved_at, created_at) VALUES (?, ?, ?, ?, ?, ?, ?)",
                ("FRZ_test", "测试冻结", "approved", json.dumps(payload), "tester", "2026-01-01", "2026-01-01"),
            )
        save_source_citation_metadata(self.project, "SRC_book", {
            "author": "张三", "title": "测试书", "place": "西安", "publisher": "测试出版社",
            "year": "1908", "type_code": "M", "verified_by": "tester",
        })
        exported = export_document(self.project, manuscript["manuscript_id"], "markdown", "builtin-tangdu-current")
        text = (self.project / exported["project_path"]).read_text(encoding="utf-8")
        self.assertEqual(text.count("## 参考文献"), 1)
        self.assertNotIn("（由工作台生成）", text)

    def test_tangdu_template_separates_references_and_explanatory_footnotes(self) -> None:
        template = next(
            item for item in ensure_journal_templates(self.project)
            if item["template_id"] == "builtin-tangdu-current"
        )
        self.assertEqual(template["requirements"]["citation_system"], "sequential_reference")
        self.assertEqual(template["requirements"]["reference_marker_position"], "superscript")
        self.assertEqual(template["requirements"]["note_role"], "explanatory_only")
        self.assertIn("2026年第2期历史论文刊例", template["version_label"])
        self.assertEqual(template["requirements"]["bibliographic_standard"], "GB/T 7714-2025")
        manuscript = import_manuscript(self.project, "唐都引注", "# 正文\n\n事实段落。")
        with self.assertRaisesRegex(ValueError, "inline sequential references"):
            create_note(
                self.project, manuscript["manuscript_id"], "NOD_missing", 4, "事实段落",
                template["template_id"], "METADATA_FIRST_PAGE_LATER", {"title": "测试书"},
            )
        with self.assertRaisesRegex(ValueError, "explicitly identified explanatory text"):
            create_note(
                self.project, manuscript["manuscript_id"], "NOD_missing", 4, "事实段落",
                template["template_id"], "REFORMAT_EXISTING", {"user_supplied_text": "说明文字。"},
            )

    def test_tangdu_direct_source_citation_requires_human_verified_original_page(self) -> None:
        manuscript = import_manuscript(
            self.project, "学术史引证", "# 正文\n\n已有研究指出这一差异。[CITE:SRC_study@PAGE_study]",
        )
        with connect(self.project) as connection:
            project_id = connection.execute("SELECT project_id FROM projects").fetchone()[0]
            connection.execute(
                "INSERT INTO sources VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
                ("SRC_study", project_id, "研究论文", "local_file", "study.pdf", "acquired", "ready", "partial", "2026-01-01"),
            )
            connection.execute(
                """INSERT INTO pages(page_id, source_id, physical_page, printed_page, page_type,
                   verification_state, use_state, machine_payload_json, human_payload_json)
                   VALUES (?, ?, ?, ?, 'content', 'human_verified', 'research_usable', '{}', '{}')""",
                ("PAGE_study", "SRC_study", 7, "23"),
            )
        save_source_citation_metadata(self.project, "SRC_study", {
            "author": "李四", "title": "秦岭考察研究", "place": "", "publisher": "",
            "year": "2024", "type_code": "J", "journal": "唐都学刊", "issue": "2",
            "page_range": "20-30", "verified_by": "tester",
        })
        exported = export_document(
            self.project, manuscript["manuscript_id"], "markdown", "builtin-tangdu-current",
        )
        text = (self.project / exported["project_path"]).read_text(encoding="utf-8")
        self.assertIn("已有研究指出这一差异。[1]23", text)
        self.assertIn("[1] 李四. 秦岭考察研究[J]. 唐都学刊，2024(2)：20-30.", text)

    def test_export_preview_warns_when_research_process_labels_enter_manuscript(self) -> None:
        manuscript = import_manuscript(
            self.project, "内部标签稿", "# 正文\n\n| 人物 | 在本文中的作用 |\n| --- | --- |\n| 甲 | 核心个案之一 |",
        )
        preview = preview_document_export(
            self.project, manuscript["manuscript_id"], "builtin-tangdu-current",
        )
        self.assertTrue(any("正文仍含研究过程语言" in item for item in preview["warnings"]))

    def test_tangdu_direct_source_citation_accepts_human_spot_checked_page(self) -> None:
        manuscript = import_manuscript(
            self.project, "学术史抽查页引证", "# 正文\n\n已有研究指出这一差异。[CITE:SRC_study@PAGE_study]",
        )
        with connect(self.project) as connection:
            project_id = connection.execute("SELECT project_id FROM projects").fetchone()[0]
            connection.execute(
                "INSERT INTO sources VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
                ("SRC_study", project_id, "研究论文", "local_file", "study.pdf", "acquired", "ready", "partial", "2026-01-01"),
            )
            connection.execute(
                """INSERT INTO pages(page_id, source_id, physical_page, printed_page, page_type,
                   verification_state, use_state, machine_payload_json, human_payload_json)
                   VALUES (?, ?, ?, ?, 'content', 'human_spot_checked', 'research_usable', '{}', '{}')""",
                ("PAGE_study", "SRC_study", 7, "23"),
            )
        save_source_citation_metadata(self.project, "SRC_study", {
            "author": "李四", "title": "秦岭考察研究", "place": "", "publisher": "",
            "year": "2024", "type_code": "J", "journal": "唐都学刊", "issue": "2",
            "page_range": "20-30", "verified_by": "tester",
        })
        exported = export_document(
            self.project, manuscript["manuscript_id"], "markdown", "builtin-tangdu-current",
        )
        text = (self.project / exported["project_path"]).read_text(encoding="utf-8")
        self.assertIn("已有研究指出这一差异。[1]23", text)
        self.assertNotIn("尚未完成逐页人工核验", "\n".join(exported["fidelity"]["warnings"]))

    def test_tangdu_superscript_range_stops_before_following_chinese_prose(self) -> None:
        from research_workbench.document_model import SEQUENTIAL_CITATION_RE

        for text, expected in (
            ("山路艰险。[5]235—237这使秦岭成为交通对象。", "[5]235—237"),
            ("前人已有讨论。[2]28李蕾、沈弘继续指出。", "[2]28"),
            ("目录见卷首。[1]I这意味着材料需要重读。", "[1]I"),
        ):
            with self.subTest(text=text):
                match = SEQUENTIAL_CITATION_RE.search(text)
                self.assertIsNotNone(match)
                self.assertEqual(match.group(0), expected)

    def test_ui_has_four_permanent_workspaces_and_nested_repair(self) -> None:
        html = (Path(__file__).parents[1] / "src" / "research_workbench" / "web_assets" / "index.html").read_text(encoding="utf-8")
        for label in ("研究对话", "研究图书馆", "文章工作台", "项目设置"):
            self.assertIn(label, html)
        self.assertIn("研究者意图基线", html)
        app = (Path(__file__).parents[1] / "src" / "research_workbench" / "web_assets" / "app.js").read_text(encoding="utf-8")
        self.assertIn("研究认知轨迹", app)
        self.assertIn("不是心理画像或来源证据", app)
        self.assertIn("$('pageJump').value = page?.physical_page || '';", app)
        self.assertIn("sessionStorage.getItem('hrwManuscriptId')", app)
        self.assertIn("sessionStorage.setItem('hrwSectionId', state.sectionId)", app)
        self.assertIn("本节尚有 ${pending.length} 份待审提案", app)
        self.assertIn("字符预算未通过", app)
        self.assertNotIn('id="sourceMode"', html)
        self.assertIn('id="openSourceRepair"', html)
        self.assertIn('id="documentCanvas"', html)
        self.assertIn('id="browserWorkbench"', html)
        self.assertIn('id="controlledBrowser"', html)
        self.assertIn("受控浏览器已打开", app)
        self.assertIn("插入正文引证并保存新修订", app)
        self.assertIn("按双栏阅读顺序重排", app)
        self.assertIn("解释性脚注文字（不得填写书目引证）", app)
        self.assertIn("保存本稿投稿信息", app)
        self.assertIn("四层分开，不自动沉淀", app)
        self.assertIn("不会自动写入外部长期记忆库", app)

    def test_submission_profile_is_scoped_to_one_manuscript(self) -> None:
        saved = save_submission_profile(self.project, self.manuscript["manuscript_id"], {
            "name": "某作者", "affiliation": "某大学", "phone": "由作者填写",
            "postal_address": "由作者填写", "postal_code": "000000",
        })
        self.assertEqual(saved["name"], "某作者")
        profiles = submission_profiles(self.project)
        self.assertEqual(profiles[0]["manuscript_id"], self.manuscript["manuscript_id"])
        self.assertEqual(profiles[0]["affiliation"], "某大学")


if __name__ == "__main__":
    unittest.main()
