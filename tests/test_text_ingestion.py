from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from research_workbench.agent_runtime import _search_source_blocks
from research_workbench.service import initialize_project, list_sources, register_source, source_view
from research_workbench.text_ingestion import ingest_docx_locator


class DocxLocatorTests(unittest.TestCase):
    def test_docx_is_searchable_but_never_evidence_qualified(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            project = root / "project"
            initialize_project(project, "DOCX locator")
            path = root / "translation.docx"
            document = Document()
            document.add_heading("秦岭行程", level=1)
            document.add_paragraph("修订地名：Han-tchong 对应汉中府。")
            document.add_paragraph("脚夫由教民村落组织。")
            document.save(path)

            source = register_source(project, path, "谭卫道日记中文修订稿")
            result = ingest_docx_locator(project, source["source_id"])
            self.assertEqual(result["status"], "locator_only")
            self.assertEqual(result["paragraph_count"], 3)

            listed = next(item for item in list_sources(project) if item["source_id"] == source["source_id"])
            self.assertEqual(listed["use_state"], "locator_only")
            self.assertGreater(listed["page_count"], 0)
            view = source_view(project, source["source_id"])
            self.assertEqual(view["pages"][0]["page_type"], "docx_locator")
            self.assertEqual(view["pages"][0]["use_state"], "locator_only")
            hit = _search_source_blocks(project, "Han-tchong/脚夫")
            self.assertEqual(len(hit), 2)
            self.assertTrue(all(item["block_use_state"] == "locator_only" for item in hit))


if __name__ == "__main__":
    unittest.main()
