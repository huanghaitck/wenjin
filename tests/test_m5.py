from __future__ import annotations

import shutil
import tempfile
import json
import threading
import time
import unittest
from pathlib import Path
from urllib.request import Request, urlopen

import pymupdf as fitz

from research_workbench.library import (
    _bibliographic_identifiers,
    _author_names,
    _filename_bibliography,
    _filename_is_identifier,
    _clean_author,
    _material_type,
    _pdf_bibliography,
    approve_candidates,
    archive_uploaded_file,
    decide_literature_relation,
    link_work_to_project,
    library_graph,
    move_work_to_shelf,
    register_author_alias,
    scan_directory,
    scan_session,
    search_library,
    update_work,
    work_detail,
)
from research_workbench.library_store import connect_library
from research_workbench.db import connect
from research_workbench.project_library import add_library_file_to_project
from research_workbench.service import initialize_project
from research_workbench.skill_registry import discover_skills
from research_workbench.web import build_server


class M5ResearchLibraryTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary = tempfile.TemporaryDirectory()
        root = Path(self.temporary.name)
        self.project = root / "project"
        self.library = root / "library"
        self.materials = root / "materials"
        self.materials.mkdir()
        initialize_project(self.project, "M5 test project")

    def tearDown(self) -> None:
        self.temporary.cleanup()

    def test_automatic_scan_skips_bench_and_transient_markdown(self) -> None:
        (self.materials / "HistRA-Bench_questions.md").write_text("benchmark", encoding="utf-8")
        (self.materials / "file.md").write_text("temporary", encoding="utf-8")
        document = fitz.open(); document.new_page().insert_text((72, 72), "Historical source")
        document.save(self.materials / "real-source.pdf"); document.close()
        scanned = scan_directory(self.project, self.materials, self.library)
        self.assertEqual([Path(item["path"]).name for item in scanned["candidates"]], ["real-source.pdf"])

    def _approve_one(self, path: Path) -> tuple[str, dict[str, object]]:
        session = scan_directory(self.project, path.parent, self.library)
        candidate = next(item for item in session["candidates"] if item["path"] == str(path.resolve()))
        approved = approve_candidates(self.project, session["session_id"], [candidate["candidate_id"]], self.library)
        return approved["approved"][0]["work_id"], candidate

    def test_one_character_change_creates_linked_version_not_new_work(self) -> None:
        source = self.materials / "chronicle.txt"
        original_text = "地方历史档案与帝国考察记录。" * 20
        source.write_text(original_text, encoding="utf-8")
        work_id, first = self._approve_one(source)

        source.write_text(source.read_text(encoding="utf-8") + "字", encoding="utf-8")
        changed = work_detail(self.project, work_id, self.library)
        self.assertEqual(changed["files"][0]["file_state"], "changed_since_last_scan")
        self.assertFalse(changed["files"][0]["versions"][0]["bytes_available"])
        session = scan_directory(self.project, self.materials, self.library)
        candidate = next(item for item in session["candidates"] if item["path"] == str(source.resolve()))
        self.assertEqual(candidate["proposed_action"], "new_version")
        self.assertEqual(candidate["existing_work_id"], work_id)
        approve_candidates(self.project, session["session_id"], [candidate["candidate_id"]], self.library)

        detail = work_detail(self.project, work_id, self.library)
        self.assertEqual(detail["file_count"], 1)
        self.assertEqual(detail["version_count"], 2)
        versions = detail["files"][0]["versions"]
        self.assertEqual(sum(item["is_current"] for item in versions), 1)
        self.assertNotEqual(first["sha256"], versions[0]["sha256"])

        source.write_text(original_text, encoding="utf-8")
        reverted = scan_directory(self.project, self.materials, self.library)
        candidate = reverted["candidates"][0]
        self.assertEqual(candidate["proposed_action"], "new_version")
        approve_candidates(self.project, reverted["session_id"], [candidate["candidate_id"]], self.library)
        self.assertEqual(work_detail(self.project, work_id, self.library)["version_count"], 3)

    def test_exact_duplicate_keeps_both_locations_and_original_bytes(self) -> None:
        original = self.materials / "source.md"
        original.write_text("# 蒙古考察史\n\n历史档案与帝国知识。" * 20, encoding="utf-8")
        before = original.read_bytes()
        work_id, _ = self._approve_one(original)
        duplicate = self.materials / "copy.md"
        shutil.copy2(original, duplicate)

        session = scan_directory(self.project, self.materials, self.library)
        candidate = next(item for item in session["candidates"] if item["path"] == str(duplicate.resolve()))
        self.assertEqual(candidate["proposed_action"], "exact_duplicate")
        approve_candidates(self.project, session["session_id"], [candidate["candidate_id"]], self.library)

        detail = work_detail(self.project, work_id, self.library)
        self.assertEqual(detail["file_count"], 2)
        self.assertEqual({item["path"] for item in detail["files"]}, {str(original.resolve()), str(duplicate.resolve())})
        self.assertEqual(original.read_bytes(), before)

    def test_pdf_triage_stops_at_ten_pages_and_library_is_searchable(self) -> None:
        pdf = self.materials / "history.pdf"
        document = fitz.open()
        for index in range(12):
            page = document.new_page()
            page.insert_text((72, 72), f"Historical archive empire chronicle page {index + 1}")
        document.set_metadata({"title": "Imperial Archive", "author": "A Historian"})
        document.save(pdf)
        document.close()

        work_id, candidate = self._approve_one(pdf)
        self.assertEqual(candidate["inspected_pages"], 10)
        self.assertEqual(candidate["page_count"], 12)
        self.assertEqual(candidate["triage_state"], "likely_historical")
        self.assertEqual(candidate["format"], "pdf")
        self.assertEqual(candidate["suggested_title"], "history")
        self.assertEqual(candidate["suggested_publisher"], "")

        updated = update_work(
            self.project,
            work_id,
            {
                "canonical_title": "Imperial Archive Revised", "author": "Professor A",
                "edition_label": "Second edition", "publication_year": "1908",
            },
            ["蒙古", "知识史"],
            self.library,
        )
        self.assertEqual(updated["canonical_title"], "Imperial Archive Revised")
        self.assertEqual(updated["editions"][0]["publication_year"], "1908")
        self.assertEqual(search_library(self.project, "Professor A", library_root=self.library)[0]["work_id"], work_id)
        self.assertEqual(search_library(self.project, "知识", library_root=self.library)[0]["work_id"], work_id)
        self.assertEqual(search_library(self.project, tags=["知识史"], library_root=self.library)[0]["work_id"], work_id)
        moved = move_work_to_shelf(self.project, work_id, "monographs", self.library)
        self.assertEqual(moved["shelf"], "monographs")
        self.assertEqual(moved["shelf_label"], "学术专著")
        graph = library_graph(self.project, library_root=self.library)
        labels = {node["label"] for node in graph["nodes"]}
        self.assertIn("Imperial Archive Revised", labels)
        self.assertEqual({node["node_type"] for node in graph["nodes"]}, {"work"})
        self.assertEqual(graph["nodes"][0]["graph_category"], "monographs")
        self.assertEqual(graph["edges"], [])
        self.assertEqual(len(graph["work_cards"]), 1)
        self.assertIn("Historical archive empire chronicle page 8", graph["work_cards"][0]["content_excerpt"])
        self.assertEqual(graph["work_cards"][0]["preview_pages"], 10)
        self.assertIsNone(graph["work_cards"][0]["project_source"])
        content_graph = library_graph(self.project, "chronicle page 8", library_root=self.library)
        self.assertEqual(content_graph["work_cards"][0]["work_id"], work_id)
        with connect_library(self.library) as connection:
            work_node = connection.execute(
                "SELECT node_id FROM knowledge_nodes WHERE node_type = 'work' AND normalized_label = ?",
                (work_id.casefold(),),
            ).fetchone()["node_id"]
            connection.execute(
                "DELETE FROM knowledge_edges WHERE source_node_id = ? OR target_node_id = ?",
                (work_node, work_node),
            )
            connection.execute("DELETE FROM knowledge_nodes WHERE node_id = ?", (work_node,))
        repaired_graph = library_graph(self.project, library_root=self.library)
        self.assertEqual(repaired_graph["backfilled_work_count"], 1)
        self.assertIn("Imperial Archive Revised", {node["label"] for node in repaired_graph["nodes"]})
        focused_graph = library_graph(self.project, limit=1, library_root=self.library)
        self.assertEqual(sum(node["node_type"] == "work" for node in focused_graph["nodes"]), 1)
        linked = link_work_to_project(self.project, work_id, self.library)
        self.assertEqual(len(linked["project_links"]), 1)
        added = add_library_file_to_project(
            self.project, self.library, work_id, linked["files"][0]["file_id"]
        )
        with connect(self.project) as connection:
            connection.execute(
                "UPDATE pages SET verification_state='human_verified', use_state='research_usable' WHERE source_id=?",
                (added["source"]["source_id"],),
            )
            connection.execute(
                """UPDATE blocks SET verification_state='human_verified', use_state='research_usable'
                   WHERE page_id IN (SELECT page_id FROM pages WHERE source_id=?)""",
                (added["source"]["source_id"],),
            )
        linked_graph = library_graph(self.project, "Imperial Archive", library_root=self.library)
        self.assertEqual(linked_graph["work_cards"][0]["project_source"]["source_id"], added["source"]["source_id"])
        content_graph = library_graph(self.project, library_root=self.library)["content_graph"]
        content_types = {node["node_type"] for node in content_graph["nodes"]}
        self.assertIn("source", content_types)
        self.assertIn("page", content_types)
        self.assertIn("content", content_types)
        self.assertTrue(any("Historical archive" in node.get("excerpt", "") for node in content_graph["nodes"]))
        self.assertIn("contains_content", {edge["relation"] for edge in content_graph["edges"]})
        version = linked["files"][0]["versions"][0]
        self.assertEqual(version["skill_name"], "historical-material-intake")
        self.assertEqual(len(version["skill_sha256"]), 64)

    def test_skill_discovery_and_uncertain_image_candidate_remain_visible(self) -> None:
        skill = next(item for item in discover_skills() if item["name"] == "historical-material-intake")
        self.assertEqual(skill["execution"], "instructions_only")
        self.assertEqual(skill["placement"], "user_action")
        self.assertIn("library_intake", skill["compatible_actions"])
        self.assertEqual(len(skill["sha256"]), 64)
        unsupported = self.materials / "scan.jpg"
        unsupported.write_bytes(b"not-a-real-image")
        session = scan_directory(self.project, self.materials, self.library)
        candidate = session["candidates"][0]
        self.assertEqual(candidate["triage_state"], "uncertain")
        approved = approve_candidates(
            self.project, session["session_id"], [candidate["candidate_id"]], self.library
        )
        self.assertEqual(len(approved["approved"]), 1)

    def test_title_page_suggestions_remain_human_editable_metadata(self) -> None:
        sample = "书 名廿二史考异\n(清)钱大昕撰\n凤凰出版社\n版次 2008年1月第1版"
        self.assertEqual(
            _pdf_bibliography(sample, "qdx", "", "", ""),
            ("廿二史考异", "(清)钱大昕撰", "凤凰出版社", "2008"),
        )

    def test_journal_first_page_precedes_filename_hints(self) -> None:
        sample = """中国北方兴隆遗址早期农业人群对木材的利用和管理
沈慧1,3, 邱振威2, 赵克良1,3, 周新郢1,3, 李小强1,3*
摘要
中文引用格式:
沈慧, 邱振威, 赵克良, 周新郢, 李小强. 2024. 中国北方兴隆遗址早期农业人群对木材的利用和管理. 中国科学: 地球科学, 54(6): 1937–1949
"""
        self.assertEqual(
            _pdf_bibliography(
                sample, "中国北方兴隆遗址早期农业人群对木材的利用和管理 沈慧",
                "", "错误标签出版社", "1888",
                path=Path("中国北方兴隆遗址早期农业人群对木材的利用和管理_沈慧.pdf"),
                material_type="article",
            ),
            (
                "中国北方兴隆遗址早期农业人群对木材的利用和管理 沈慧",
                "沈慧；邱振威；赵克良；周新郢；李小强",
                "中国科学: 地球科学",
                "2024",
            ),
        )

    def test_journal_filename_author_is_used_only_when_confirmed_on_page(self) -> None:
        sample = """农业考古2021·5
俄罗斯学者视野下的近代中俄茶叶贸易与晋商
刘啸虎
李 珂
摘要：晋商与茶叶贸易研究。
"""
        result = _pdf_bibliography(
            sample, "俄罗斯学者视野下的近代中俄茶叶贸易与晋商 刘啸虎", "", "", "",
            path=Path("俄罗斯学者视野下的近代中俄茶叶贸易与晋商_刘啸虎.pdf"),
            material_type="article",
        )
        self.assertEqual(result[1], "刘啸虎")
        self.assertIn("农业考古", result[2])

    def test_article_signals_precede_archival_words_in_title(self) -> None:
        self.assertEqual(
            _material_type("档案史料研究", "某大学学报 2024年第2期 文章编号"),
            "article",
        )

    def test_article_prose_does_not_masquerade_as_a_journal_title(self) -> None:
        sample = """历史地图中的解释
成一农
摘要：本文由中国社会科学院历史研究所历史地理研究室与复旦大学历史地理研究中心共同讨论。
"""
        result = _pdf_bibliography(
            sample, "历史地图中的解释 成一农", "", "", "",
            path=Path("历史地图中的解释_成一农.pdf"), material_type="article",
        )
        self.assertEqual(result[1], "成一农")
        self.assertEqual(result[2], "")

    def test_english_article_reads_author_journal_and_year_from_first_page(self) -> None:
        sample = """Shen Hou
Nature's Tonic: Beer,
Ecology, and Urbanization
in a Chinese City, 1900-50
Shen Hou, “Nature's Tonic: Beer, Ecology, and Urbanization in a Chinese City, 1900-50,” Environmental History 24 (2019): 282-306
"""
        result = _pdf_bibliography(
            sample, "Nature's Tonic: Beer, Ecology, and Urbanization in a Chinese City, 1900-50",
            "", "", "", material_type="article",
        )
        self.assertEqual(result[1], "Shen Hou")
        self.assertEqual(result[2], "Environmental History")
        self.assertEqual(result[3], "2019")

    def test_cyrillic_title_page_responsibility_is_read(self) -> None:
        sample = """Путешествие по Китаю в 1874 - 1875 гг.
Из дневника члена экспедиции
П. Я. Пясецкого.
ТОМ II.
МОСКВА. 1882.
"""
        result = _pdf_bibliography(sample, "Puteshestvie", "", "", "", material_type="book_or_document")
        self.assertEqual(result[1], "П. Я. Пясецкого")
        self.assertEqual(_clean_author(result[1]), "П. Я. Пясецкого")

    def test_multi_author_metadata_connects_each_shared_author(self) -> None:
        self.assertEqual(_author_names("甲；乙"), ["甲", "乙"])
        first = self.materials / "first-author.md"; second = self.materials / "second-author.md"
        first.write_text("# 第一篇\n\n历史材料。" * 20, encoding="utf-8")
        second.write_text("# 第二篇\n\n历史材料。" * 20, encoding="utf-8")
        first_work, _ = self._approve_one(first); second_work, _ = self._approve_one(second)
        update_work(self.project, first_work, {"canonical_title": "第一篇", "author": "甲；乙"}, [], self.library)
        update_work(self.project, second_work, {"canonical_title": "第二篇", "author": "甲"}, [], self.library)
        graph = library_graph(self.project, library_root=self.library)
        self.assertTrue(any(edge["relation"] == "same_author" for edge in graph["edges"]))
        self.assertTrue({"甲", "乙"}.issubset({node["label"] for node in graph["entity_nodes"] if node["node_type"] == "person"}))
        self.assertGreaterEqual(sum(edge["relation"] == "authored_by" for edge in graph["entity_edges"]), 3)

    def test_author_aliases_preserve_raw_names_but_share_one_entity(self) -> None:
        first = self.materials / "english-name.md"; second = self.materials / "chinese-name.md"
        first.write_text("# English work\n\nHistory." * 20, encoding="utf-8")
        second.write_text("# 中文作品\n\n历史。" * 20, encoding="utf-8")
        first_work, _ = self._approve_one(first); second_work, _ = self._approve_one(second)
        update_work(self.project, first_work, {"canonical_title": "English work", "author": "Shen Hou"}, [], self.library)
        update_work(self.project, second_work, {"canonical_title": "中文作品", "author": "侯深"}, [], self.library)
        register_author_alias(self.library, "Shen Hou", "侯深", "tester", "verified bilingual byline")
        graph = library_graph(self.project, library_root=self.library)
        self.assertTrue(any(edge["relation"] == "same_author" for edge in graph["edges"]))
        people = [node["label"] for node in graph["entity_nodes"] if node["node_type"] == "person"]
        self.assertIn("侯深", people); self.assertNotIn("Shen Hou", people)
        self.assertFalse(any(node["node_type"] == "tag" and node["label"].startswith(("metadata:", "triage:", "material:", "shelf:")) for node in graph["entity_nodes"]))
        searched = library_graph(self.project, "侯深", library_root=self.library)
        self.assertEqual({card["work_id"] for card in searched["work_cards"]}, {first_work, second_work})

    def test_docx_uses_page_bibliography_before_filename_hints(self) -> None:
        from docx import Document
        from research_workbench.library import _inspect_file

        path = self.materials / "期刊文章_副本.docx"
        document = Document()
        document.add_paragraph("地方史研究中的材料问题")
        document.add_paragraph("王小明")
        document.add_paragraph("某大学学报 2023年第2期")
        document.add_paragraph("摘要：本文讨论地方史材料。")
        document.save(path)
        inspected = _inspect_file(path)
        self.assertEqual(inspected["suggested_author"], "王小明")
        self.assertEqual(inspected["suggested_publisher"], "某大学学报")
        self.assertEqual(inspected["suggested_year"], "2023")

    def test_bibliographic_identifiers_merge_same_scan_files_and_refresh_cards(self) -> None:
        self.assertEqual(
            _bibliographic_identifiers("DOI: 10.4000/books.pul.13146 ISBN: 9782729711061"),
            {"doi:10.4000/books.pul.13146", "isbn:9782729711061"},
        )
        page = (
            "Regional History (1245-1800)\nNetworks and Perceptions\n"
            "Jean-Louis Gaulin and Susanne Rau (dir.)\n"
            "DOI: 10.4000/books.pul.13146\nISBN: 9782729711061\n"
        )
        for name, suffix in (("Regional History.pdf", "Original"), ("Regional History(1).pdf", "Generated copy")):
            document = fitz.open()
            document.new_page().insert_textbox(fitz.Rect(72, 72, 520, 760), page + suffix)
            document.save(self.materials / name)
            document.close()
        scanned = scan_directory(self.project, self.materials, self.library)
        actions = {Path(item["path"]).name: item["proposed_action"] for item in scanned["candidates"]}
        self.assertEqual(actions["Regional History.pdf"], "register_new")
        self.assertEqual(actions["Regional History(1).pdf"], "same_scan_work")
        approved = approve_candidates(self.project, scanned["session_id"], None, self.library)
        self.assertEqual(len({item["work_id"] for item in approved["approved"]}), 1)
        refreshed = scan_session(self.project, scanned["session_id"], self.library)
        self.assertTrue(all(item["resolved_file_count"] == 2 for item in refreshed["candidates"]))
        self.assertTrue(all(item["resolved_work_author"] == "Jean-Louis Gaulin；Susanne Rau" for item in refreshed["candidates"]))
        self.assertTrue(all(item["resolved_shelf"] == "monographs" for item in refreshed["candidates"]))

    def test_bulk_word_gate_keeps_research_files_and_explicit_upload_can_override(self) -> None:
        from docx import Document

        samples = {
            "摘要.docx": "摘要：这是孤立片段。",
            "研究生外出申请审批表.docx": "申请人和审批意见。",
            "旅行日记人工翻译稿.docx": "人工翻译的旅行日记与史料。",
            "葡萄根瘤蚜相关材料.docx": "农业报刊中的葡萄材料。",
            "环境史读书报告.docx": "读书报告与文献评论。",
            "~$临时.docx": "Office 临时文件。",
            "编者前言.docx": "孤立的译稿前言。",
        }
        for name, text in samples.items():
            document = Document(); document.add_paragraph(text * 20); document.save(self.materials / name)
        scanned = scan_directory(self.project, self.materials, self.library)
        self.assertEqual(scanned["total_count"], 3)
        self.assertEqual(scanned["ignored_word_count"], 4)
        self.assertEqual(scanned["word_review_count"], 1)
        self.assertEqual(
            {Path(item["path"]).name for item in scanned["candidates"]},
            {"旅行日记人工翻译稿.docx", "葡萄根瘤蚜相关材料.docx", "环境史读书报告.docx"},
        )
        bulk = approve_candidates(self.project, scanned["session_id"], None, self.library)
        self.assertEqual(len(bulk["approved"]), 2)
        uploaded = archive_uploaded_file(
            self.project, self.materials / "研究生外出申请审批表.docx", self.library,
        )
        self.assertTrue(uploaded["work_id"])

    def test_file_name_is_the_intake_title_and_copy_suffixes_are_cleaned(self) -> None:
        self.assertEqual(
            _filename_bibliography(Path("清代西北史料_12940532 (2).pdf")),
            ("清代西北史料", "", ""),
        )
        self.assertEqual(
            _filename_bibliography(Path("秦岭葡萄（农业未来报1883年5月13日）.docx")),
            ("秦岭葡萄（农业未来报1883年5月13日）", "1883", "农业未来报"),
        )
        self.assertEqual(
            _filename_bibliography(Path("20241358_研究者_谭卫道秦岭考察(1).docx"))[0],
            "研究者 谭卫道秦岭考察",
        )
        self.assertTrue(_filename_is_identifier("2604.24690v1"))
        self.assertFalse(_filename_is_identifier("Catholic Missionaries on the Shu Roads"))
        self.assertEqual(_clean_author("CNKI"), "")
        self.assertEqual(_clean_author("Tiziana Lioi"), "Tiziana Lioi")

    def test_exact_registered_title_in_markdown_creates_a_traceable_relation(self) -> None:
        target = self.materials / "target.md"
        target.write_text("# Referenced Monograph\n\nHistorical methods and archives." * 10, encoding="utf-8")
        target_work, _ = self._approve_one(target)
        update_work(self.project, target_work, {"canonical_title": "Referenced Monograph"}, [], self.library)
        source = self.materials / "source-study.pdf"
        document = fitz.open()
        document.new_page().insert_text((72, 72), "Source Study main argument and evidence.")
        document.new_page().insert_text((72, 72), "References: Referenced Monograph. University Press.")
        document.set_metadata({"title": "Source Study", "author": "Scholar B"})
        document.save(source)
        document.close()
        source_work, _ = self._approve_one(source)
        detail = update_work(self.project, source_work, {"canonical_title": "Source Study"}, [], self.library)
        add_library_file_to_project(self.project, self.library, source_work, detail["files"][0]["file_id"])
        graph = library_graph(self.project, library_root=self.library)
        candidate = next(item for item in graph["literature_relations"] if item["target_work_id"] == target_work)
        self.assertEqual(candidate["status"], "derived")
        self.assertIn(candidate["relation_type"], {edge["relation"] for edge in graph["edges"]})
        decided = decide_literature_relation(
            self.project, candidate["relation_key"], False, candidate["relation_type"], "researcher", "false match",
            self.library,
        )
        self.assertEqual(decided["status"], "rejected")
        refreshed = library_graph(self.project, library_root=self.library)
        rejected = next(item for item in refreshed["literature_relations"] if item["relation_key"] == candidate["relation_key"])
        self.assertEqual(rejected["status"], "rejected")

    def test_scan_suggests_library_shelves_and_bulk_approval_keeps_files_in_place(self) -> None:
        samples = {
            "source.md": ("# 地方志史料\n\n地方志日记与档案。", "primary_sources"),
            "article.md": ("# 区域史研究论文\n\n某大学学报期刊论文。", "academic_articles"),
            "book.md": ("# 区域史专著\n\n某某出版社 ISBN 978-7。", "monographs"),
            "draft.md": ("# 我的返修稿\n\n尚未刊行的论文稿。", "personal_manuscripts"),
            "notes.md": ("# 环境史读书笔记\n\n本周阅读札记。", "reading_notes"),
            "catalog.md": ("# 地方文献目录索引\n\n工具书与目录。", "reference_works"),
            "unknown.md": ("# 普通材料\n\n内容尚待研究者判断。", "unclassified"),
        }
        original = {}
        for name, (content, _) in samples.items():
            path = self.materials / name
            path.write_text(content * 20, encoding="utf-8")
            original[path] = path.read_bytes()
        session = scan_directory(self.project, self.materials, self.library)
        shelves = {Path(item["path"]).name: item["suggested_shelf"] for item in session["candidates"]}
        self.assertEqual(shelves, {name: expected for name, (_, expected) in samples.items()})
        self.assertEqual(session["eligible_remaining_count"], len(samples))
        result = approve_candidates(self.project, session["session_id"], None, self.library)
        self.assertEqual(len(result["approved"]), len(samples))
        self.assertTrue(all(path.read_bytes() == value for path, value in original.items()))
        works = search_library(self.project, library_root=self.library)
        self.assertEqual({work["shelf"] for work in works}, set(shelves.values()))
        graph = library_graph(self.project, library_root=self.library)
        self.assertNotIn("notes", {node["label"] for node in graph["nodes"] if node["node_type"] == "work"})
        self.assertEqual({node["node_type"] for node in graph["nodes"]}, {"work"})
        self.assertIn("academic_articles", {node.get("graph_category") for node in graph["nodes"]})

    def test_same_clean_title_registers_one_work_with_multiple_files(self) -> None:
        first = self.materials / "first"
        second = self.materials / "second"
        first.mkdir(); second.mkdir()
        (first / "Shared historical study.md").write_text("# A\n\nHistory archive material." * 20, encoding="utf-8")
        (second / "Shared historical study.md").write_text("# B\n\nHistory archive material revised." * 20, encoding="utf-8")
        session = scan_directory(self.project, self.materials, self.library)
        approved = approve_candidates(self.project, session["session_id"], None, self.library)
        self.assertEqual(len(approved["approved"]), 2)
        work_ids = {item["work_id"] for item in approved["approved"]}
        self.assertEqual(len(work_ids), 1)
        detail = work_detail(self.project, work_ids.pop(), self.library)
        self.assertEqual(detail["file_count"], 2)

    def test_material_type_uses_project_location_only_as_a_classification_hint(self) -> None:
        self.assertEqual(_material_type("未定名", "内容", r"D:\\研究\\个人论文与稿件\\草稿.docx"), "personal_manuscript")
        self.assertEqual(_material_type("某书", "CIP 某出版社"), "monograph")

    def test_loopback_library_api_scans_approves_and_shows_versions(self) -> None:
        source = self.materials / "api-history.txt"
        source.write_text("历史档案与地方志研究。" * 30, encoding="utf-8")
        server = build_server(self.project, port=0, library_root=self.library)
        worker = threading.Thread(target=server.serve_forever, daemon=True)
        worker.start()
        base = f"http://127.0.0.1:{server.server_port}"

        def post(path: str, payload: dict[str, object]) -> dict[str, object]:
            request = Request(
                base + path,
                data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
                headers={"Content-Type": "application/json"},
                method="POST",
            )
            return json.loads(urlopen(request, timeout=5).read())

        try:
            snapshot = json.loads(urlopen(base + "/api/snapshot", timeout=5).read())
            self.assertEqual(snapshot["library"]["library_root"], str(self.library.resolve()))
            session = post("/api/library/scan", {"source_root": str(self.materials)})
            deadline = time.monotonic() + 5
            while session["status"] == "scanning" and time.monotonic() < deadline:
                time.sleep(0.02)
                session = json.loads(urlopen(
                    base + f"/api/library/scan?id={session['session_id']}", timeout=5
                ).read())
            candidate = session["candidates"][0]
            result = post(
                "/api/library/approve",
                {"session_id": session["session_id"], "candidate_ids": [candidate["candidate_id"]]},
            )
            work_id = result["approved"][0]["work_id"]
            detail = json.loads(urlopen(base + f"/api/library/work?id={work_id}", timeout=5).read())
            self.assertEqual(detail["version_count"], 1)
            self.assertEqual(len(detail["files"][0]["versions"][0]["sha256"]), 64)
            original = urlopen(
                base + f"/api/library/file?id={detail['files'][0]['file_id']}", timeout=5
            ).read()
            self.assertEqual(original, source.read_bytes())
        finally:
            server.shutdown()
            server.server_close()
            worker.join(timeout=5)


if __name__ == "__main__":
    unittest.main()
