from __future__ import annotations

import hashlib
import io
import json
import re
import uuid
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.shared import RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

from .authoring import ensure_journal_templates
from .citations import check_note_anchors, list_notes
from .db import append_audit, connect, utc_now
from .docx_notes import add_footnote_reference, attach_footnotes
from .readiness import formal_research_readiness


def _id(prefix: str) -> str:
    return f"{prefix}_{uuid.uuid4().hex}"


def _json(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False, sort_keys=True)


def _table_cells(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _is_table_divider(line: str) -> bool:
    cells = _table_cells(line)
    return len(cells) >= 2 and all(cell.replace(":", "").replace("-", "") == "" and "-" in cell for cell in cells)


def _nodes_from_text(text: str) -> list[dict[str, Any]]:
    lines = text.replace("\r\n", "\n").splitlines()
    nodes: list[dict[str, Any]] = []
    paragraph: list[str] = []

    def flush_paragraph() -> None:
        value = "\n".join(paragraph).strip()
        if value:
            nodes.append({"type": "paragraph", "node_id": _id("NOD"), "text": value})
        paragraph.clear()

    index = 0
    while index < len(lines):
        line = lines[index]
        if line.startswith("### "):
            flush_paragraph()
            nodes.append({"type": "subheading", "node_id": _id("NOD"), "text": line[4:].strip()})
            index += 1
            continue
        if "|" in line and index + 1 < len(lines) and _is_table_divider(lines[index + 1]):
            flush_paragraph()
            rows = [_table_cells(line)]
            index += 2
            while index < len(lines) and "|" in lines[index] and lines[index].strip():
                rows.append(_table_cells(lines[index]))
                index += 1
            width = max(len(row) for row in rows)
            rows = [row + [""] * (width - len(row)) for row in rows]
            nodes.append({"type": "table", "node_id": _id("NOD"), "rows": rows})
            continue
        if not line.strip():
            flush_paragraph()
        else:
            paragraph.append(line)
        index += 1
    flush_paragraph()
    return nodes or [
        {"type": "paragraph", "node_id": _id("NOD"), "text": ""}
    ]


def _table_markdown(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    lines = ["| " + " | ".join(row) + " |" for row in normalized]
    lines.insert(1, "| " + " | ".join("---" for _ in range(width)) + " |")
    return "\n".join(lines)


def _node_text(node: dict[str, Any]) -> str:
    if node.get("type") == "table":
        return _table_markdown([[str(cell) for cell in row] for row in node.get("rows", [])])
    return str(node.get("text", ""))


def _tree_from_sections(title: str, sections: list[dict[str, Any]]) -> dict[str, Any]:
    return {
        "type": "document",
        "node_id": _id("NOD"),
        "title": title,
        "children": [
            {
                "type": "section",
                "node_id": _id("NOD"),
                "section_id": section["section_id"],
                "heading": section["heading"],
                "children": _nodes_from_text(section["content"]),
            }
            for section in sections
        ],
    }


def _plain_text(tree: dict[str, Any]) -> str:
    lines = [str(tree.get("title", ""))]
    for section in tree.get("children", []):
        lines.append(str(section.get("heading", "")))
        lines.extend(_node_text(node) for node in section.get("children", []))
    return "\n".join(lines)


def _export_sections(tree: dict[str, Any]) -> list[dict[str, Any]]:
    """Hide an empty legacy title section; the document title is already rendered separately."""
    title = str(tree.get("title", "")).strip()
    return [
        section for section in tree.get("children", [])
        if not (
            str(section.get("heading", "")).strip() == title
            and not any(_node_text(node).strip() for node in section.get("children", []))
        )
    ]


def _validate_tree(tree: dict[str, Any]) -> None:
    if tree.get("type") != "document" or not isinstance(tree.get("children"), list):
        raise ValueError("document must be a structured document tree")
    seen: set[str] = set()
    for section in tree["children"]:
        if section.get("type") != "section" or not str(section.get("heading", "")).strip():
            raise ValueError("each document section requires a heading")
        if not isinstance(section.get("children"), list):
            raise ValueError("each document section requires paragraph nodes")
        for node in [section, *section["children"]]:
            node_id = str(node.get("node_id", ""))
            if not node_id or node_id in seen:
                raise ValueError("document node IDs must be present and unique")
            seen.add(node_id)
        if any(node.get("type") not in {"paragraph", "quote", "list_item", "subheading", "table"} for node in section["children"]):
            raise ValueError("unsupported document node type")
        for node in section["children"]:
            if node.get("type") != "table":
                continue
            rows = node.get("rows")
            if not isinstance(rows, list) or not rows or any(not isinstance(row, list) for row in rows):
                raise ValueError("table nodes require rows")
            width = len(rows[0])
            if width < 2 or any(len(row) != width for row in rows):
                raise ValueError("table rows must have the same width and at least two columns")
            if any(not isinstance(cell, str) for row in rows for cell in row):
                raise ValueError("table cells must contain text")


def ensure_document(project_root: Path, manuscript_id: str) -> dict[str, Any]:
    with connect(project_root) as connection:
        manuscript = connection.execute(
            "SELECT * FROM manuscripts WHERE manuscript_id = ?", (manuscript_id,)
        ).fetchone()
        if manuscript is None:
            raise KeyError(f"unknown manuscript: {manuscript_id}")
        existing = connection.execute(
            "SELECT document_id FROM manuscript_documents WHERE manuscript_id = ?", (manuscript_id,)
        ).fetchone()
        if existing is None:
            sections = [dict(row) for row in connection.execute(
                """SELECT s.section_id, s.heading, v.content
                   FROM manuscript_sections s JOIN section_versions v ON v.version_id = s.current_version_id
                   WHERE s.manuscript_id = ? ORDER BY s.section_order""", (manuscript_id,)
            )]
            tree = _tree_from_sections(str(manuscript["title"]), sections)
            document_id, revision_id, now = _id("DOC"), _id("DREV"), utc_now()
            digest = hashlib.sha256(_plain_text(tree).encode("utf-8")).hexdigest()
            fidelity = {"level": "exact_text", "warnings": [], "source": "legacy_sections"}
            connection.execute(
                "INSERT INTO manuscript_documents(document_id, manuscript_id, current_revision_id, created_at, updated_at) VALUES (?, ?, ?, ?, ?)",
                (document_id, manuscript_id, revision_id, now, now),
            )
            connection.execute(
                """INSERT INTO document_revisions(revision_id, document_id, base_revision_id, document_json,
                   plain_text_hash, source_format, status, fidelity_json, created_at)
                   VALUES (?, ?, NULL, ?, ?, 'legacy_sections', 'approved', ?, ?)""",
                (revision_id, document_id, _json(tree), digest, _json(fidelity), now),
            )
    return document_detail(project_root, manuscript_id)


def document_detail(project_root: Path, manuscript_id: str) -> dict[str, Any]:
    with connect(project_root) as connection:
        row = connection.execute(
            """SELECT d.*, r.document_json, r.plain_text_hash, r.source_format, r.status,
                      r.fidelity_json, r.created_at AS revision_created_at
               FROM manuscript_documents d JOIN document_revisions r ON r.revision_id = d.current_revision_id
               WHERE d.manuscript_id = ?""", (manuscript_id,)
        ).fetchone()
        if row is None:
            raise KeyError(f"structured document is not initialized: {manuscript_id}")
        revisions = [dict(item) for item in connection.execute(
            """SELECT revision_id, base_revision_id, plain_text_hash, source_format, status,
                      fidelity_json, created_at FROM document_revisions
               WHERE document_id = ? ORDER BY created_at DESC, revision_id DESC""", (row["document_id"],)
        )]
        receipts = [dict(item) for item in connection.execute(
            "SELECT * FROM document_io_receipts WHERE manuscript_id = ? ORDER BY created_at DESC",
            (manuscript_id,),
        )]
    result = dict(row)
    result["document"] = json.loads(result.pop("document_json"))
    result["fidelity"] = json.loads(result.pop("fidelity_json"))
    for item in revisions:
        item["fidelity"] = json.loads(item.pop("fidelity_json"))
    for item in receipts:
        item["fidelity"] = json.loads(item.pop("fidelity_json"))
    result["revisions"], result["io_receipts"] = revisions, receipts
    result["notes"] = list_notes(project_root, manuscript_id)
    return result


def save_document(project_root: Path, manuscript_id: str, tree: dict[str, Any],
                  source_format: str = "structured_editor",
                  fidelity: dict[str, Any] | None = None) -> dict[str, Any]:
    _validate_tree(tree)
    current = ensure_document(project_root, manuscript_id)
    revision_id, now = _id("DREV"), utc_now()
    fidelity = fidelity or {"level": "native", "warnings": [], "source": "structured_editor"}
    with connect(project_root) as connection:
        existing = {row["section_id"]: dict(row) for row in connection.execute(
            """SELECT s.*, v.content AS current_content,
                      v.evidence_refs_json AS current_evidence_refs_json
               FROM manuscript_sections s
               LEFT JOIN section_versions v ON v.version_id = s.current_version_id
               WHERE s.manuscript_id = ?""", (manuscript_id,)
        )}
        for section in tree["children"]:
            if str(section.get("section_id", "")) not in existing:
                section["section_id"] = _id("SEC")
        digest = hashlib.sha256(_plain_text(tree).encode("utf-8")).hexdigest()
        connection.execute(
            """INSERT INTO document_revisions(revision_id, document_id, base_revision_id, document_json,
               plain_text_hash, source_format, status, fidelity_json, created_at)
               VALUES (?, ?, ?, ?, ?, ?, 'approved', ?, ?)""",
            (revision_id, current["document_id"], current["current_revision_id"], _json(tree), digest,
             source_format, _json(fidelity), now),
        )
        connection.execute(
            "UPDATE manuscript_sections SET section_order = -section_order WHERE manuscript_id = ?",
            (manuscript_id,),
        )
        for order, section in enumerate(tree["children"], start=1):
            section_id = str(section.get("section_id", ""))
            if section_id not in existing:
                connection.execute(
                    "INSERT INTO manuscript_sections(section_id, manuscript_id, section_order, heading, created_at) VALUES (?, ?, ?, ?, ?)",
                    (section_id, manuscript_id, order, str(section["heading"]).strip(), now),
                )
            content = "\n\n".join(_node_text(node).strip() for node in section["children"]).strip()
            prior = existing.get(section_id, {})
            base = prior.get("current_version_id")
            if not content and str(prior.get("current_content") or "").strip():
                raise ValueError(
                    "a non-empty manuscript section cannot be saved as blank; reload and retry"
                )
            version_id = base
            if not prior or content != str(prior.get("current_content") or ""):
                version_id = _id("SEV")
                evidence_refs = _json(list(dict.fromkeys(re.findall(r"\[EVID:([A-Za-z0-9_]+)\]", content))))
                connection.execute(
                    """INSERT INTO section_versions(version_id, section_id, base_version_id, operation, content,
                       evidence_refs_json, model_snapshot_json, status, created_at, approved_at)
                       VALUES (?, ?, ?, 'manual_structured_edit', ?, ?, '{"provider":"human_editor"}', 'approved', ?, ?)""",
                    (version_id, section_id, base, content, evidence_refs, now, now),
                )
            connection.execute(
                "UPDATE manuscript_sections SET section_order = ?, heading = ?, current_version_id = ? WHERE section_id = ?",
                (order, str(section["heading"]).strip(), version_id, section_id),
            )
        connection.execute(
            "UPDATE manuscript_documents SET current_revision_id = ?, updated_at = ? WHERE document_id = ?",
            (revision_id, now, current["document_id"]),
        )
        connection.execute("UPDATE manuscripts SET title = ?, updated_at = ? WHERE manuscript_id = ?",
                           (str(tree.get("title", "")).strip() or "未命名稿件", now, manuscript_id))
        append_audit(connection, "document_revision_saved", "manuscript", manuscript_id,
                     {"revision_id": revision_id, "base_revision_id": current["current_revision_id"]})
    check_note_anchors(project_root, manuscript_id, tree)
    return document_detail(project_root, manuscript_id)


def sync_approved_section(project_root: Path, manuscript_id: str, section_id: str,
                          content: str, section_version_id: str) -> dict[str, Any]:
    """Make an approved section version the current structured-document revision."""
    current = ensure_document(project_root, manuscript_id)
    tree = deepcopy(current["document"])
    with connect(project_root) as connection:
        latest = connection.execute(
            "SELECT heading, current_version_id FROM manuscript_sections "
            "WHERE manuscript_id = ? AND section_id = ?",
            (manuscript_id, section_id),
        ).fetchone()
    if latest is None:
        raise KeyError(f"unknown manuscript section: {section_id}")
    if latest["current_version_id"] != section_version_id:
        raise ValueError("section version is stale; reload the manuscript before synchronizing")
    section = next((item for item in tree["children"] if item.get("section_id") == section_id), None)
    if section is None:
        raise KeyError(f"structured document is missing section: {section_id}")
    old_heading = str(section.get("heading", ""))
    section["heading"] = str(latest["heading"])
    old_children = section.get("children", [])
    new_children = _nodes_from_text(content)
    for index, node in enumerate(new_children):
        if index < len(old_children):
            node["node_id"] = old_children[index]["node_id"]
    same_content = [_node_text(node) for node in old_children] == [_node_text(node) for node in new_children]
    if same_content and old_heading == str(latest["heading"]):
        return current
    section["children"] = new_children
    revision_id, now = _id("DREV"), utc_now()
    digest = hashlib.sha256(_plain_text(tree).encode("utf-8")).hexdigest()
    fidelity = {
        "level": "evidence_reviewed", "warnings": [],
        "source": "approved_section_version", "section_version_id": section_version_id,
    }
    with connect(project_root) as connection:
        connection.execute(
            """INSERT INTO document_revisions(revision_id, document_id, base_revision_id, document_json,
               plain_text_hash, source_format, status, fidelity_json, created_at)
               VALUES (?, ?, ?, ?, ?, 'approved_section_version', 'approved', ?, ?)""",
            (revision_id, current["document_id"], current["current_revision_id"], _json(tree), digest,
             _json(fidelity), now),
        )
        connection.execute(
            "UPDATE manuscript_documents SET current_revision_id = ?, updated_at = ? WHERE document_id = ?",
            (revision_id, now, current["document_id"]),
        )
        append_audit(connection, "approved_section_synced", "manuscript", manuscript_id, {
            "section_id": section_id, "section_version_id": section_version_id,
            "revision_id": revision_id, "base_revision_id": current["current_revision_id"],
        })
    check_note_anchors(project_root, manuscript_id, tree)
    return document_detail(project_root, manuscript_id)


def _numbered_notes(notes: list[dict[str, Any]], tree: dict[str, Any]) -> tuple[dict[str, list[tuple[int, dict[str, Any]]]], list[dict[str, Any]]]:
    by_node: dict[str, list[tuple[int, dict[str, Any]]]] = {}
    node_order = {
        str(node.get("node_id")): position
        for position, node in enumerate(
            node for section in tree.get("children", []) for node in section.get("children", [])
        )
    }
    ordered = sorted(notes, key=lambda item: (
        node_order.get(str(item["anchor_node_id"]), 10**9), int(item["anchor_offset"]), item["created_at"]
    ))
    for number, note in enumerate(ordered, start=1):
        by_node.setdefault(str(note["anchor_node_id"]), []).append((number, note))
    return by_node, ordered


def _markdown_text(text: str, placements: list[tuple[int, dict[str, Any]]]) -> str:
    for number, note in sorted(placements, key=lambda item: int(item[1]["anchor_offset"]), reverse=True):
        offset = min(max(0, int(note["anchor_offset"])), len(text))
        text = text[:offset] + f"[^note{number}]" + text[offset:]
    return text


def markdown_from_tree(tree: dict[str, Any], notes: list[dict[str, Any]] | None = None) -> str:
    by_node, ordered = _numbered_notes(notes or [], tree)
    lines = [f"# {tree.get('title', '未命名稿件')}", ""]
    for section in _export_sections(tree):
        lines.extend([f"## {section.get('heading', '正文')}", ""])
        for node in section.get("children", []):
            if node.get("type") == "table":
                lines.extend([_table_markdown(node.get("rows", [])), ""])
                continue
            prefix = "### " if node.get("type") == "subheading" else ("> " if node.get("type") == "quote" else "")
            text = _markdown_text(str(node.get("text", "")), by_node.get(str(node.get("node_id")), []))
            lines.extend([prefix + text, ""])
    if ordered:
        lines.append("---")
        lines.append("")
        for number, note in enumerate(ordered, start=1):
            lines.append(f"[^note{number}]: {note['rendered_text']}")
    return "\n".join(lines).rstrip() + "\n"


def _docx_blocks(document: DocxDocument) -> list[Paragraph | Table]:
    blocks: list[Paragraph | Table] = []
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            blocks.append(Paragraph(child, document))
        elif child.tag == qn("w:tbl"):
            blocks.append(Table(child, document))
    return blocks


def _docx_table_rows(table: Table) -> list[list[str]]:
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    width = max((len(row) for row in rows), default=0)
    return [row + [""] * (width - len(row)) for row in rows]


EVIDENCE_MARKER_RE = re.compile(r"(?:\[EVID:[A-Za-z0-9_]+\])+")
DIRECT_CITATION_MARKER_RE = re.compile(r"\[CITE:([A-Za-z0-9_]+)@([A-Za-z0-9_:]+)\]")
REFERENCE_MARKER_RE = re.compile(
    r"(?:\[EVID:[A-Za-z0-9_]+\])+|\[CITE:[A-Za-z0-9_]+@[A-Za-z0-9_:]+\]"
)
SEQUENTIAL_CITATION_RE = re.compile(
    r"(\[\d+\](?:原书页待核|[0-9IVXLCDMivxlcdm]+(?:[—–\-、,，][0-9IVXLCDMivxlcdm]+)*))"
)


def _add_text_runs(paragraph: Any, text: str, superscript_citations: bool) -> None:
    def add_fragment(fragment: str) -> None:
        cursor = 0
        for match in re.finditer(r"\*\*(.+?)\*\*", fragment):
            paragraph.add_run(fragment[cursor:match.start()])
            run = paragraph.add_run(match.group(1))
            run.bold = True
            cursor = match.end()
        paragraph.add_run(fragment[cursor:])

    if not superscript_citations:
        add_fragment(text)
        return
    cursor = 0
    for match in SEQUENTIAL_CITATION_RE.finditer(text):
        add_fragment(text[cursor:match.start()])
        run = paragraph.add_run(match.group(0))
        run.font.superscript = True
        cursor = match.end()
    add_fragment(text[cursor:])


def _reference_entry(number: int, metadata: dict[str, Any]) -> str:
    author, title = metadata.get("author", ""), metadata.get("title", "")
    type_code, year = metadata.get("type_code", ""), metadata.get("year", "")
    page_range = metadata.get("page_range", "")
    if type_code == "J":
        issue = metadata.get("issue", "")
        volume = str(metadata.get("volume", ""))
        if volume:
            year_issue = f"{year},{volume}{f'({issue})' if issue else ''}"
        else:
            year_issue = f"{year}{f'({issue})' if issue else ''}"
        journal_tail = "，".join(filter(None, (metadata.get("journal", ""), year_issue)))
        return f"[{number}] {author}. {title}[J]. {journal_tail}{('：' + page_range) if page_range else ''}."
    responsibility = f"，{metadata['translator']}译" if metadata.get("translator") else ""
    edition = f"，{metadata['edition']}" if metadata.get("edition") else ""
    publication = "：".join(value for value in (metadata.get("place", ""), metadata.get("publisher", "")) if value)
    tail = "，".join(value for value in (publication, year) if value)
    return f"[{number}] {author}. {title}{responsibility}{edition}[{type_code}]. {tail}{('：' + page_range) if page_range else ''}.".strip()


def _prepare_sequential_references(project_root: Path, tree: dict[str, Any],
                                   manuscript_id: str = "") -> tuple[dict[str, Any], list[str], str]:
    prepared = deepcopy(tree)
    evidence: dict[str, dict[str, Any]] = {}
    with connect(project_root) as connection:
        for row in connection.execute(
            "SELECT payload_json FROM evidence_freezes WHERE status = 'approved' ORDER BY created_at DESC"
        ):
            payload = json.loads(row["payload_json"])
            for claim in payload.get("claims", []):
                for item in claim.get("evidence", []):
                    evidence.setdefault(str(item.get("evidence_id", "")), item)
        metadata = {
            row["source_id"]: dict(row)
            for row in connection.execute("SELECT * FROM source_citation_metadata")
        }
        pages = {
            (row["source_id"], row["page_id"]): dict(row)
            for row in connection.execute(
                "SELECT source_id, page_id, physical_page, printed_page, verification_state, use_state FROM pages"
            )
        }
        page_printed = {page_id: str(row["printed_page"] or "").strip()
                        for (_, page_id), row in pages.items()}
        profile_row = connection.execute(
            "SELECT profile_json FROM manuscript_submission_profiles WHERE manuscript_id = ?",
            (manuscript_id,),
        ).fetchone() if manuscript_id else None
        submission_profile = json.loads(profile_row["profile_json"]) if profile_row else {}
    warnings: list[str] = []
    source_numbers: dict[str, int] = {}

    def marker(match: re.Match[str]) -> str:
        direct = DIRECT_CITATION_MARKER_RE.fullmatch(match.group(0))
        if direct:
            source_id, page_id = direct.groups()
            page = pages.get((source_id, page_id))
            if page is None:
                warnings.append(f"来源引证 {source_id} 的原页 {page_id} 不存在")
                return match.group(0)
            if page["use_state"] != "research_usable" or page["verification_state"] not in {
                "human_spot_checked", "human_verified", "human_repaired",
            }:
                warnings.append(f"来源引证 {source_id} 的原页 {page_id} 尚未完成逐页人工核验")
                return match.group(0)
            printed_page = str(page["printed_page"] or "").strip()
            if not printed_page:
                warnings.append(
                    f"来源引证 {source_id} 的物理页 {page['physical_page']} 尚未确认原书页码"
                )
                return match.group(0)
            if source_id not in source_numbers:
                source_numbers[source_id] = len(source_numbers) + 1
            return f"[{source_numbers[source_id]}]{printed_page}"
        grouped: dict[str, list[dict[str, Any]]] = {}
        for evidence_id in re.findall(r"\[EVID:([A-Za-z0-9_]+)\]", match.group(0)):
            item = evidence.get(evidence_id)
            if item is None:
                warnings.append(f"证据 {evidence_id} 未出现在已批准冻结包中")
                continue
            grouped.setdefault(str(item["source_id"]), []).append(item)
        rendered: list[str] = []
        for source_id, items in grouped.items():
            if source_id not in source_numbers:
                source_numbers[source_id] = len(source_numbers) + 1
            printed_pages = list(dict.fromkeys(
                str(page)
                for item in items
                for page in (
                    item.get("printed_pages", [])
                    or [page_printed.get(str(page_id), "") for page_id in item.get("page_ids", [])]
                )
                if str(page).strip()
            ))
            if not printed_pages:
                digital = list(dict.fromkeys(
                    str(page) for item in items for page in item.get("physical_pages", []) if str(page).strip()
                ))
                warnings.append(f"来源 {source_id} 缺少原书页码；数字页 {','.join(digital) or '未知'} 仅供回查")
                page_text = "原书页待核"
            else:
                page_text = "、".join(printed_pages)
            rendered.append(f"[{source_numbers[source_id]}]{page_text}")
        return "；".join(rendered) if rendered else match.group(0)

    for section in prepared.get("children", []):
        for node in section.get("children", []):
            if node.get("type") == "table":
                node["rows"] = [[REFERENCE_MARKER_RE.sub(marker, str(cell)) for cell in row]
                                for row in node.get("rows", [])]
            else:
                node["text"] = REFERENCE_MARKER_RE.sub(marker, str(node.get("text", "")))

    references: list[str] = []
    for source_id, number in sorted(source_numbers.items(), key=lambda item: item[1]):
        item = metadata.get(source_id)
        if not item or item.get("verification_status") != "HUMAN_VERIFIED":
            warnings.append(f"来源 {source_id} 尚未人工核验引文元数据")
            continue
        required = ["author", "title", "year", "type_code"]
        if item.get("type_code") == "M":
            required.extend(["place", "publisher"])
        elif item.get("type_code") == "J":
            required.append("journal")
        missing = [field for field in required if not str(item.get(field, "")).strip()]
        if missing:
            warnings.append(f"来源 {source_id} 缺少引文元数据：{','.join(missing)}")
        references.append(_reference_entry(number, item))
    if source_numbers:
        prepared["children"] = [
            section for section in prepared["children"]
            if not (
                str(section.get("heading", "")).strip() == "参考文献"
                and all(
                    not re.match(r"^\s*\[\d+\]", _node_text(node))
                    for node in section.get("children", [])
                )
            )
        ]
        reference_section = {
            "type": "section", "node_id": _id("NOD"), "section_id": "",
            "heading": "参考文献", "children": [
                {"type": "paragraph", "node_id": _id("NOD"), "text": entry}
                for entry in references
            ],
        }
        insertion = next((index for index, section in enumerate(prepared["children"])
                          if any(term in str(section.get("heading", "")) for term in ("英文", "English", "作者简介", "联系方式"))),
                         len(prepared["children"]))
        prepared["children"].insert(insertion, reference_section)
    headings = [str(section.get("heading", "")) for section in prepared.get("children", [])]
    if len(str(prepared.get("title", ""))) > 20:
        warnings.append("《唐都学刊》题名一般不超过20个汉字；当前题名需人工复核")
    if len(_plain_text(prepared).replace(" ", "").replace("\n", "")) < 10000:
        warnings.append("《唐都学刊》来稿约10000字并欢迎万字以上；当前正文不足10000字符")
    if not any("摘要" in heading for heading in headings):
        warnings.append("缺少约300字中文摘要")
    if not any("关键词" in heading for heading in headings):
        warnings.append("缺少3—8个中文关键词")
    if not any("英文" in heading or "English" in heading for heading in headings):
        warnings.append("《唐都学刊》要求参考文献后附英文题名、作者、单位、摘要和关键词")
    if not any("作者" in heading or "联系方式" in heading for heading in headings):
        warnings.append("正文结构中未见作者简介或投稿信息段；可在期刊模板页填写独立投稿信息")
    missing_profile = [label for key, label in (
        ("name", "姓名"), ("affiliation", "工作单位"), ("phone", "联系电话"),
        ("postal_address", "详细邮寄地址"), ("postal_code", "邮编"),
    ) if not submission_profile.get(key)]
    if missing_profile:
        warnings.append("当前稿件投稿信息缺少：" + "、".join(missing_profile))
    elif not any("投稿信息" in heading for heading in headings):
        profile_lines = [
            "；".join(filter(None, (
                f"姓名：{submission_profile.get('name', '')}",
                f"真实姓名：{submission_profile.get('real_name', '')}" if submission_profile.get("real_name") else "",
                f"性别：{submission_profile.get('gender', '')}" if submission_profile.get("gender") else "",
                f"民族：{submission_profile.get('ethnicity', '')}" if submission_profile.get("ethnicity") else "",
                f"籍贯：{submission_profile.get('native_place', '')}" if submission_profile.get("native_place") else "",
            ))),
            "；".join(filter(None, (
                f"工作单位：{submission_profile.get('affiliation', '')}",
                f"学位及学科：{submission_profile.get('degree', '')}" if submission_profile.get("degree") else "",
                f"职称：{submission_profile.get('professional_title', '')}" if submission_profile.get("professional_title") else "",
                f"职务：{submission_profile.get('position', '')}" if submission_profile.get("position") else "",
            ))),
            f"主要研究方向：{submission_profile.get('research_interests', '')}" if submission_profile.get("research_interests") else "",
            "；".join(filter(None, (
                f"项目来源：{submission_profile.get('project_source', '')}" if submission_profile.get("project_source") else "",
                f"项目编号：{submission_profile.get('project_number', '')}" if submission_profile.get("project_number") else "",
            ))),
            "；".join(filter(None, (
                f"联系电话：{submission_profile['phone']}",
                f"详细邮寄地址：{submission_profile['postal_address']}",
                f"邮编：{submission_profile['postal_code']}",
                f"电子邮箱：{submission_profile.get('email', '')}" if submission_profile.get("email") else "",
            ))),
        ]
        prepared["children"].append({
            "type": "section", "node_id": _id("NOD"), "section_id": "", "heading": "作者投稿信息",
            "children": [{"type": "paragraph", "node_id": _id("NOD"), "text": line}
                         for line in profile_lines if line],
        })
    exported_text = _plain_text(prepared)
    if "待作者填写" in exported_text or "to be supplied by the researcher" in exported_text:
        warnings.append("英文作者单位及作者投稿信息仍为人工占位，提交前必须由作者填写")
    return prepared, list(dict.fromkeys(warnings)), "READY" if not warnings else "BLOCKED"


def preview_document_export(project_root: Path, manuscript_id: str, template_id: str) -> dict[str, Any]:
    detail = ensure_document(project_root, manuscript_id)
    templates = {item["template_id"]: item for item in ensure_journal_templates(project_root)}
    if template_id not in templates:
        raise KeyError(f"unknown journal template: {template_id}")
    template = templates[template_id]
    export_tree = detail["document"]
    warnings: list[str] = []
    citation_status = "NOT_APPLICABLE"
    if "参考文献置于文后" in str(template.get("citation_style", "")):
        export_tree, warnings, citation_status = _prepare_sequential_references(
            project_root, export_tree, manuscript_id,
        )
    work_language = sorted(set(re.findall(
        r"在本文中的作用|核心个案(?:之一)?|核心窗口|时间锚|观察段|正文时段",
        _plain_text(export_tree),
    )))
    if work_language:
        warnings.append("正文仍含研究过程语言：" + "、".join(work_language))
    readiness = formal_research_readiness(project_root)
    if readiness["status"] != "READY":
        warnings.extend(f"正式研究门禁：{item}" for item in readiness["blockers"])
    notes = list_notes(project_root, manuscript_id, approved_only=True)
    return {
        "manuscript_id": manuscript_id,
        "revision_id": detail["current_revision_id"],
        "template": template,
        "document": export_tree,
        "markdown": markdown_from_tree(export_tree, notes),
        "notes": notes,
        "warnings": warnings,
        "citation_status": citation_status,
        "formal_research_status": readiness["status"],
        "research_readiness": readiness,
    }


def import_docx(project_root: Path, title: str, data: bytes) -> dict[str, Any]:
    from .authoring import import_manuscript

    document = Document(io.BytesIO(data))
    lines: list[str] = []
    for block in _docx_blocks(document):
        if isinstance(block, Table):
            rows = _docx_table_rows(block)
            if rows and len(rows[0]) >= 2:
                lines.extend([_table_markdown(rows), ""])
            continue
        text = block.text.strip()
        if not text:
            continue
        style = block.style.name if block.style else ""
        if style.startswith("Heading"):
            level = style.removeprefix("Heading").strip() or "1"
            lines.append("#" * min(6, max(1, int(level))) + " " + text)
        else:
            lines.extend([text, ""])
    if not lines:
        raise ValueError("DOCX does not contain readable paragraphs")
    manuscript = import_manuscript(project_root, title.strip() or "导入的 DOCX", "\n".join(lines))
    detail = ensure_document(project_root, manuscript["manuscript_id"])
    warnings = []
    if document.tables:
        warnings.append(f"{len(document.tables)} 个表格已导入为可编辑结构；复杂合并单元格需人工复核")
    warnings.append("批注、修订、域代码、脚注与嵌入对象未做无损往返保证")
    _record_receipt(project_root, manuscript["manuscript_id"], detail["current_revision_id"],
                    "import", "docx", "", {"level": "limited", "warnings": warnings})
    result = document_detail(project_root, manuscript["manuscript_id"])
    result["import_fidelity"] = {"level": "limited", "warnings": warnings}
    return result


def reimport_docx(project_root: Path, manuscript_id: str, data: bytes) -> dict[str, Any]:
    current = ensure_document(project_root, manuscript_id)
    document = Document(io.BytesIO(data))
    old_sections = current["document"].get("children", [])
    sections: list[dict[str, Any]] = []
    active: dict[str, Any] | None = None
    for block in _docx_blocks(document):
        if isinstance(block, Table):
            rows = _docx_table_rows(block)
            if not rows or len(rows[0]) < 2:
                continue
            if active is None:
                active = {
                    "type": "section", "node_id": _id("NOD"), "section_id": "",
                    "heading": "正文", "children": [],
                }
                sections.append(active)
            active["children"].append({"type": "table", "node_id": _id("NOD"), "rows": rows})
            continue
        text = block.text.strip()
        if not text:
            continue
        style = block.style.name if block.style else ""
        if style.startswith("Heading"):
            active = {
                "type": "section", "node_id": _id("NOD"), "section_id": "",
                "heading": text, "children": [],
            }
            sections.append(active)
            continue
        if active is None:
            active = {
                "type": "section", "node_id": _id("NOD"), "section_id": "",
                "heading": "正文", "children": [],
            }
            sections.append(active)
        node_type = "quote" if style == "Quote" else ("list_item" if style.startswith("List") else "paragraph")
        active["children"].append({"type": node_type, "node_id": _id("NOD"), "text": text})
    if not sections:
        raise ValueError("DOCX does not contain readable paragraphs")
    for index, section in enumerate(sections):
        if index < len(old_sections):
            section["section_id"] = str(old_sections[index].get("section_id", ""))
        if not section["children"]:
            section["children"].append({"type": "paragraph", "node_id": _id("NOD"), "text": ""})
    warnings = ["Word 修改稿已作为新修订导回，原导出与旧修订保持不变",
                "批注、修订、域代码、脚注与嵌入对象未做无损往返保证",
                "原结构化注释锚点因 Word 段落身份变化而需要重新核对"]
    if document.tables:
        warnings.append(f"{len(document.tables)} 个表格已导回为可编辑结构；复杂合并单元格需人工复核")
    tree = {
        "type": "document", "node_id": _id("NOD"),
        "title": str(current["document"].get("title", "未命名稿件")), "children": sections,
    }
    result = save_document(
        project_root, manuscript_id, tree, "docx_reimport",
        {"level": "limited", "warnings": warnings, "source": "microsoft_word_roundtrip"},
    )
    _record_receipt(project_root, manuscript_id, result["current_revision_id"], "import", "docx", "",
                    {"level": "limited", "warnings": warnings})
    result = document_detail(project_root, manuscript_id)
    result["import_fidelity"] = {"level": "limited", "warnings": warnings}
    return result


def export_document(project_root: Path, manuscript_id: str, format_name: str,
                    template_id: str = "builtin-history-research") -> dict[str, Any]:
    preview = preview_document_export(project_root, manuscript_id, template_id)
    detail = {"current_revision_id": preview["revision_id"]}
    template = preview["template"]
    export_tree = preview["document"]
    warnings = list(preview["warnings"])
    citation_status = preview["citation_status"]
    notes = preview["notes"]
    by_node, ordered_notes = _numbered_notes(notes, export_tree)
    export_root = project_root / "exports"
    export_root.mkdir(parents=True, exist_ok=True)
    if format_name == "markdown":
        path = export_root / f"{manuscript_id}-{detail['current_revision_id']}.md"
        path.write_text(preview["markdown"], encoding="utf-8")
        level = "native_text"
    elif format_name == "docx":
        path = export_root / f"{manuscript_id}-{detail['current_revision_id']}.docx"
        document = Document()
        requirements = template.get("requirements", {})
        superscript_citations = requirements.get("reference_marker_position") == "superscript"
        section_properties = document.sections[0]
        section_properties.page_width, section_properties.page_height = Cm(21), Cm(29.7)
        normal = document.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        normal.font.size = Pt(float(requirements.get("body_size_pt", 12)))
        for style_name, size in (("Heading 1", 14), ("Heading 2", 12)):
            style = document.styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
            style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor(0, 0, 0)
        title = document.add_paragraph()
        title_run = title.add_run(str(export_tree.get("title", "未命名稿件")))
        title_run.bold = True
        title_run.font.name = "Times New Roman"
        title_run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        title_run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        title_run.font.size = Pt(18)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for section in _export_sections(export_tree):
            document.add_heading(str(section.get("heading", "正文")), level=1)
            for node in section.get("children", []):
                if node.get("type") == "table":
                    rows = node.get("rows", [])
                    if rows:
                        table = document.add_table(rows=len(rows), cols=len(rows[0]))
                        table.style = "Table Grid"
                        table.autofit = False
                        total_width = 15.2
                        weights = [0.16, 0.18, 0.24, 0.21, 0.21] if len(rows[0]) == 5 else [1 / len(rows[0])] * len(rows[0])
                        for row_index, row in enumerate(rows):
                            properties = table.rows[row_index]._tr.get_or_add_trPr()
                            properties.append(OxmlElement("w:cantSplit"))
                            if row_index == 0:
                                properties.append(OxmlElement("w:tblHeader"))
                            for column_index, value in enumerate(row):
                                cell = table.cell(row_index, column_index)
                                cell.width = Cm(total_width * weights[column_index])
                                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                cell.text = str(value)
                                if row_index == 0:
                                    for run in cell.paragraphs[0].runs:
                                        run.bold = True
                    continue
                if node.get("type") == "subheading":
                    paragraph = document.add_heading(str(node.get("text", "")), level=2)
                    continue
                style = "Quote" if node.get("type") == "quote" else None
                paragraph = document.add_paragraph(style=style)
                text = str(node.get("text", ""))
                cursor = 0
                for number, note in sorted(by_node.get(str(node.get("node_id")), []), key=lambda item: int(item[1]["anchor_offset"])):
                    offset = min(max(cursor, int(note["anchor_offset"])), len(text))
                    _add_text_runs(paragraph, text[cursor:offset], superscript_citations)
                    add_footnote_reference(paragraph, number)
                    cursor = offset
                _add_text_runs(paragraph, text[cursor:], superscript_citations)
        document.save(path)
        attach_footnotes(path, [note["rendered_text"] for note in ordered_notes],
                         requirements.get("number_restart") == "each_page")
        level = "structured_with_true_footnotes"
        warnings.append("已写入真实 Word 脚注；圈码外观、每页重编号及最终分页仍须在目标 Word 版本中复核")
        warnings.append("LibreOffice 无界面转换在中文标点后紧接脚注时可能失败；Microsoft Word 渲染已通过")
    else:
        raise ValueError("format must be markdown or docx")
    fidelity = {
        "level": level, "warnings": warnings, "approved_note_count": len(ordered_notes),
        "template_id": template_id, "template_revision_id": template.get("template_revision_id"),
        "template_verification_status": template.get("verification_status"),
        "citation_status": citation_status,
        "formal_research_status": preview["formal_research_status"],
    }
    receipt = _record_receipt(project_root, manuscript_id, detail["current_revision_id"],
                              "export", format_name, path.relative_to(project_root).as_posix(), fidelity)
    return {**receipt, "download_url": "/api/export/file?path=" + path.relative_to(project_root).as_posix()}


def _record_receipt(project_root: Path, manuscript_id: str, revision_id: str, direction: str,
                    format_name: str, project_path: str, fidelity: dict[str, Any]) -> dict[str, Any]:
    receipt_id, now = _id("IOR"), utc_now()
    with connect(project_root) as connection:
        connection.execute(
            """INSERT INTO document_io_receipts(receipt_id, manuscript_id, revision_id, direction,
               format, project_path, fidelity_json, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
            (receipt_id, manuscript_id, revision_id, direction, format_name, project_path, _json(fidelity), now),
        )
    return {"receipt_id": receipt_id, "manuscript_id": manuscript_id, "revision_id": revision_id,
            "direction": direction, "format": format_name, "project_path": project_path,
            "fidelity": fidelity, "created_at": now}
